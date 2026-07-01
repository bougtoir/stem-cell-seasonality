#!/usr/bin/env python3
"""
Comprehensive manuscript verification for Nature Methods Perspective.

Checks:
1. All figures cited in text, in sequential order
2. In-text citation numbers ↔ reference list cross-check (no orphans either way)
3. Citation numbers are sequential (no gaps, no exceeding total)
4. Crossref API reference existence verification
5. Word count within Nature Methods limits
"""

import os
import re
import json
import time
import urllib.request
import urllib.parse
from docx import Document

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
MANUSCRIPT = os.path.join(OUTPUT_DIR, "NatMethods_Perspective_InvisibleVariables.docx")
SCRIPT_PATH = os.path.join(os.path.dirname(__file__),
                            "create_natmethods_perspective_docx.py")


def extract_text_from_docx(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_superscript_citations(path):
    """Extract superscript citation numbers from docx runs."""
    doc = Document(path)
    citations = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.superscript:
                text = run.text.strip()
                for part in re.split(r'[,\s]+', text):
                    m = re.match(r'(\d+)[-\u2013\u2014](\d+)', part)
                    if m:
                        citations.extend(range(int(m.group(1)), int(m.group(2)) + 1))
                    elif part.isdigit():
                        citations.append(int(part))
    return citations


def get_body_text(path):
    """Extract body text between first heading and References."""
    doc = Document(path)
    body = []
    in_body = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "The methodological blind spot":
            in_body = True
            continue
        if text in ("Figure Legends", "References"):
            in_body = False
            continue
        if in_body and text:
            body.append(text)
    return "\n".join(body)


def check_figures(full_text):
    """Check figure citations appear in text in sequential order."""
    results = []
    fig_positions = {}
    for m in re.finditer(r'Figure\s+(\d+)', full_text):
        num = int(m.group(1))
        if num not in fig_positions:
            fig_positions[num] = m.start()

    expected = [1, 2]
    for n in expected:
        if n in fig_positions:
            results.append({"figure": n, "status": "CITED",
                             "position": fig_positions[n]})
        else:
            results.append({"figure": n, "status": "NOT_CITED", "position": -1})

    # Check sequential order
    positions = [fig_positions.get(n, -1) for n in expected]
    sequential = all(a < b for a, b in zip(positions, positions[1:])
                     if a >= 0 and b >= 0)

    return results, sequential


def check_citations(citations, ref_count):
    """Check citation numbers against reference list."""
    unique_cited = sorted(set(citations))
    issues = []

    # Check for citations exceeding reference count
    over = [c for c in unique_cited if c > ref_count]
    if over:
        issues.append(f"Citations exceed reference count ({ref_count}): {over}")

    # Check for uncited references
    all_refs = set(range(1, ref_count + 1))
    uncited = sorted(all_refs - set(unique_cited))
    if uncited:
        issues.append(f"Uncited references: {uncited}")

    # Check for gaps
    if unique_cited:
        expected = set(range(1, max(unique_cited) + 1))
        gaps = sorted(expected - set(unique_cited))
        if gaps:
            issues.append(f"Citation gaps: {gaps}")

    # Check first-appearance order
    first_appearance = {}
    for c in citations:
        if c not in first_appearance:
            first_appearance[c] = len(first_appearance) + 1
    out_of_order = []
    for ref_num, appearance_order in sorted(first_appearance.items()):
        if ref_num != appearance_order:
            out_of_order.append((ref_num, appearance_order))

    return {
        "unique_cited": unique_cited,
        "uncited": uncited,
        "over_range": over,
        "out_of_order": out_of_order,
        "issues": issues,
    }


def verify_reference_exists(ref_text, ref_num):
    """Query Crossref for reference existence."""
    # Skip books
    if '(Wiley,' in ref_text or 'Press,' in ref_text:
        return {"ref_num": ref_num, "status": "BOOK_SKIPPED", "detail": ref_text}

    # Nature format: Author, A. B. et al. Title. J. Abbrev. vol, pages (year).
    # Extract title: find text after "et al. " or after last "X. " (author initial)
    # then before the journal abbreviation
    # Strategy: use regex to find title between authors and journal
    title = ""
    author = ref_text.split(',')[0].strip()

    # Nature format title extraction:
    # After "et al. " or after last single-initial ". ", before journal abbreviation
    # Use greedy match for title, stop at pattern like ". Abbrev. " or ". vol,"
    m = re.search(r'(?:et al\. |[A-Z]\. )([A-Z][^.]{10,})\. ', ref_text)
    if m:
        title = m.group(1)
    else:
        # Fallback: split by '. ' and find the longest segment (likely the title)
        parts = ref_text.split('. ')
        if len(parts) >= 2:
            title = max(parts, key=len)

    query = f"{author} {title}"
    url = ("https://api.crossref.org/works?"
           + urllib.parse.urlencode({"query": query, "rows": 3}))

    try:
        req = urllib.request.Request(url, headers={
            "User-Agent": "NatMethods-Verify/1.0 (mailto:verify@example.com)"
        })
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode())

        items = data.get("message", {}).get("items", [])
        if not items:
            return {"ref_num": ref_num, "status": "NOT_FOUND", "detail": query}

        # Check title overlap across all returned items
        best_overlap = 0
        best_item = None
        ref_words = set(title.lower().split())

        for item in items:
            cr_title = item.get("title", [""])[0].lower()
            cr_words = set(cr_title.split())
            if ref_words and cr_words:
                overlap = len(ref_words & cr_words) / max(len(ref_words), 1)
                if overlap > best_overlap:
                    best_overlap = overlap
                    best_item = item

        if best_overlap >= 0.4:
            doi = best_item.get("DOI", "no-doi")
            cr_title = best_item.get("title", [""])[0]
            return {
                "ref_num": ref_num,
                "status": "VERIFIED",
                "doi": doi,
                "crossref_title": cr_title,
                "overlap": round(best_overlap, 2),
            }
        else:
            return {
                "ref_num": ref_num,
                "status": "LOW_MATCH",
                "overlap": round(best_overlap, 2),
                "detail": query,
            }

    except Exception as e:
        return {"ref_num": ref_num, "status": "ERROR", "detail": str(e)}


def get_references_from_script():
    """Import REFERENCES list from the manuscript script."""
    import importlib.util
    spec = importlib.util.spec_from_file_location("natmethods", SCRIPT_PATH)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod.REFERENCES


def main():
    print("=" * 60)
    print("Nature Methods Perspective — Manuscript Verification")
    print("=" * 60)

    if not os.path.exists(MANUSCRIPT):
        print(f"\nERROR: Manuscript not found at {MANUSCRIPT}")
        print("Run create_natmethods_perspective_docx.py first.")
        return

    full_text = extract_text_from_docx(MANUSCRIPT)
    body_text = get_body_text(MANUSCRIPT)
    refs = get_references_from_script()
    ref_count = len(refs)

    results = {
        "manuscript": os.path.basename(MANUSCRIPT),
        "reference_count": ref_count,
    }

    # 1. Word count
    body_words = len(body_text.split())
    abstract_text = ""
    in_abstract = False
    for line in full_text.split("\n"):
        if line.strip() == "Abstract":
            in_abstract = True
            continue
        if line.strip() == "The methodological blind spot":
            in_abstract = False
            continue
        if in_abstract and line.strip():
            abstract_text += line + " "
    abstract_words = len(abstract_text.split())

    print(f"\n1. WORD COUNT")
    print(f"   Body text: {body_words} words (limit: ~4,000)")
    print(f"   Abstract:  {abstract_words} words (limit: 150)")
    results["body_word_count"] = body_words
    results["abstract_word_count"] = abstract_words
    results["body_within_limit"] = body_words <= 4000
    results["abstract_within_limit"] = abstract_words <= 150

    # 2. Figure checks
    print(f"\n2. FIGURE CITATIONS")
    fig_results, sequential = check_figures(full_text)
    for fr in fig_results:
        status = "OK" if fr["status"] == "CITED" else "MISSING"
        print(f"   Figure {fr['figure']}: {status}")
    print(f"   Sequential order: {'YES' if sequential else 'NO'}")
    results["figures"] = fig_results
    results["figures_sequential"] = sequential

    # 3. Citation checks
    print(f"\n3. CITATION CROSS-CHECK")
    citations = extract_superscript_citations(MANUSCRIPT)
    cit_results = check_citations(citations, ref_count)
    print(f"   Unique citations in text: {len(cit_results['unique_cited'])}")
    print(f"   References in list: {ref_count}")
    if cit_results["uncited"]:
        print(f"   UNCITED references: {cit_results['uncited']}")
    if cit_results["over_range"]:
        print(f"   OUT OF RANGE citations: {cit_results['over_range']}")
    if cit_results["out_of_order"]:
        print(f"   OUT OF ORDER (ref#, appeared#): {cit_results['out_of_order']}")
    else:
        print(f"   Citation order: SEQUENTIAL (Vancouver compliant)")
    if not cit_results["issues"]:
        print(f"   Status: ALL OK — {ref_count} references, all cited")
    results["citations"] = cit_results

    # 4. Crossref verification
    print(f"\n4. CROSSREF REFERENCE VERIFICATION")
    ref_results = []
    for i, ref in enumerate(refs, 1):
        r = verify_reference_exists(ref, i)
        ref_results.append(r)
        status = r["status"]
        extra = f" (DOI: {r['doi']})" if status == "VERIFIED" else ""
        print(f"   [{i:2d}] {status}{extra}")
        time.sleep(0.5)  # be polite to Crossref

    verified = sum(1 for r in ref_results if r["status"] == "VERIFIED")
    books = sum(1 for r in ref_results if r["status"] == "BOOK_SKIPPED")
    issues = sum(1 for r in ref_results
                 if r["status"] not in ("VERIFIED", "BOOK_SKIPPED"))
    print(f"\n   Summary: {verified} VERIFIED, {books} BOOK, {issues} ISSUES")
    results["references"] = ref_results

    # 5. Output files check
    print(f"\n5. OUTPUT FILES")
    expected_files = [
        "NatMethods_Perspective_InvisibleVariables.docx",
        "NatMethods_Figure1.png",
        "NatMethods_Figure1.tiff",
        "NatMethods_Figure2.png",
        "NatMethods_Figure2.tiff",
        "NatMethods_Figures.pptx",
        "NatMethods_Presubmission_Inquiry.docx",
    ]
    file_results = []
    for f in expected_files:
        path = os.path.join(OUTPUT_DIR, f)
        exists = os.path.exists(path)
        size = os.path.getsize(path) if exists else 0
        print(f"   {'OK' if exists else 'MISSING'}: {f}"
              + (f" ({size:,} bytes)" if exists else ""))
        file_results.append({"file": f, "exists": exists, "size": size})
    results["files"] = file_results

    # Save reports
    json_path = os.path.join(OUTPUT_DIR, "natmethods_manuscript_verification.json")
    with open(json_path, "w") as fp:
        json.dump(results, fp, indent=2, default=str)
    print(f"\nJSON report: {json_path}")

    md_path = os.path.join(OUTPUT_DIR, "natmethods_manuscript_verification.md")
    with open(md_path, "w") as fp:
        fp.write("# Nature Methods Perspective — Manuscript Verification\n\n")
        fp.write(f"## Word Count\n")
        fp.write(f"- Body: {body_words} / 4,000\n")
        fp.write(f"- Abstract: {abstract_words} / 150\n\n")
        fp.write(f"## Figures\n")
        for fr in fig_results:
            fp.write(f"- Figure {fr['figure']}: {fr['status']}\n")
        fp.write(f"- Sequential: {'Yes' if sequential else 'No'}\n\n")
        fp.write(f"## Citations\n")
        fp.write(f"- {len(cit_results['unique_cited'])}/{ref_count} "
                 f"references cited in text\n")
        if cit_results["uncited"]:
            fp.write(f"- Uncited: {cit_results['uncited']}\n")
        if cit_results["out_of_order"]:
            fp.write(f"- Out of order: {cit_results['out_of_order']}\n")
        else:
            fp.write(f"- Order: Sequential (Vancouver compliant)\n")
        fp.write(f"\n## References (Crossref)\n")
        for r in ref_results:
            status = r["status"]
            doi = r.get("doi", "")
            fp.write(f"- [{r['ref_num']}] {status}"
                     + (f" — {doi}" if doi else "") + "\n")
        fp.write(f"\nSummary: {verified} VERIFIED, {books} BOOK, "
                 f"{issues} ISSUES\n\n")
        fp.write(f"## Output Files\n")
        for fr in file_results:
            fp.write(f"- {'OK' if fr['exists'] else 'MISSING'}: {fr['file']}\n")
    print(f"Markdown report: {md_path}")

    print(f"\n{'=' * 60}")
    all_ok = (
        results["body_within_limit"]
        and results["abstract_within_limit"]
        and sequential
        and not cit_results["issues"]
        and issues == 0
        and all(f["exists"] for f in file_results)
    )
    print(f"OVERALL: {'ALL CHECKS PASSED' if all_ok else 'ISSUES FOUND'}")
    print(f"{'=' * 60}")

    return results


if __name__ == "__main__":
    main()
