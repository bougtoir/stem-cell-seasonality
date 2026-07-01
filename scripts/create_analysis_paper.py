#!/usr/bin/env python3
"""
Generate a manuscript (docx) presenting the GEO × Solar Activity exploratory analysis
in standard IMRaD format (Introduction, Methods, Results, Discussion).

Target: Brief Communication / Research Letter accompanying the Perspective.
Format: ~2,500 words, 2 figures, 1 table, Vancouver citations.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ──────────────────────────────────────────────
# References (Vancouver, first-appearance order)
# ──────────────────────────────────────────────
REFERENCES = [
    # 1
    "Yamanaka S. Pluripotent stem cell-based cell therapy—promise and challenges. "
    "Cell Stem Cell. 2020;27(4):523–531.",
    # 2
    "Volpato V, Smith J, Bhinge A, et al. Reproducibility of molecular phenotypes after "
    "long-term differentiation to human iPSC-derived neurons: a multi-site omics study. "
    "Stem Cell Reports. 2018;11(4):897–911.",
    # 3
    "Fossati V, Bhinge A, Bhatt R, et al. Addressing variability in iPSC-derived models "
    "of human disease: guidelines to promote reproducibility. Dis Model Mech. "
    "2020;13(1):dmm042317.",
    # 4
    "Leathersich SJ, Hart RJ, Wijs LA, et al. Season at the time of oocyte collection "
    "and frozen embryo transfer outcomes. Hum Reprod. 2023;38(9):1761–1770.",
    # 5
    "Suzuki H, Togashi M, Adachi J, Toyoda Y. Seasonal variation in cell cycle during "
    "early development of the mouse embryo. Reproduction. 1992;94(2):431–436.",
    # 6
    "Umemura Y, Maki I, Tsuchiya Y, et al. The circadian clock CRY1 regulates "
    "pluripotent stem cell identity and somatic cell reprogramming. Cell Rep. "
    "2023;42(6):112590.",
    # 7
    "Sato Y, Bando H, Di Piazza M, et al. The environmental risk assessment of "
    "cell-processing facilities for cell therapy in a Japanese academic institution. "
    "PLoS One. 2020;15(8):e0236600.",
    # 8
    "Diatroptova MA, Kosyreva AM, Makarova OV, Diatroptov ME. About 4-day rhythm of "
    "proliferative activity of L-929 cells in culture correlates with the intensity of "
    "secondary cosmic radiation fluctuations. Bull Exp Biol Med. 2022;173(4):531–535.",
    # 9
    "Edgar R, Domrachev M, Lash AE. Gene Expression Omnibus: NCBI gene expression "
    "and hybridization array data repository. Nucleic Acids Res. 2002;30(1):207–210.",
    # 10
    "Hathaway DH. The solar cycle. Living Rev Sol Phys. 2015;12(1):4.",
    # 11
    "Li X, Zhang Y, Chen H, et al. Real-time monitoring reveals the effects of low "
    "concentrations of volatile organic compounds in the embryology laboratory. "
    "Reprod Biomed Online. 2025;50(2):103876.",
    # 12
    "Lapidot T, Kollet O. Daily light and darkness onset and circadian rhythms "
    "metabolically synchronize hematopoietic stem cell differentiation and maintenance. "
    "Exp Hematol. 2019;78:1–10.",
    # 13
    "Ortmann D, Vallier L. Variability of human pluripotent stem cell lines. "
    "Curr Opin Genet Dev. 2017;46:179–185.",
    # 14
    "Klein SG, Hendriks WT, Reusken C, et al. Toward best practices for controlling "
    "mammalian cell culture environments. Front Cell Dev Biol. 2022;10:788808.",
    # 15
    "Kirkeby A, Main H, Carpenter M. Pluripotent stem-cell-derived therapies in clinical "
    "trial: a 2025 update. Cell Stem Cell. 2025;32(3):329–331.",
]


def add_superscript_refs(paragraph, text):
    """Parse text with {N} or {N-M} or {N,M} markers and create superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
    return paragraph


def set_paragraph_format(para, space_after=Pt(6), space_before=Pt(0),
                         line_spacing=1.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    para.paragraph_format.line_spacing = line_spacing
    para.alignment = alignment
    return para


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_table(doc):
    """Add Table 1: Summary statistics."""
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'

    headers = ["Parameter", "Value", "Note"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    data = [
        ("Total GEO Series", "6,101", "iPSC/ESC differentiation, 2001–2026"),
        ("Records with country data", "600 (10%)", "SOFT header contact_country"),
        ("Northern Hemisphere", "584 (97%)", "USA, China, UK, Germany, Japan, etc."),
        ("Southern Hemisphere", "16 (3%)", "Australia, Brazil, South Africa"),
        ("Monthly distribution (chi-square)", "\u03c7\u00b2 = 32.3, p = 0.0007", "Reject uniform null"),
        ("Peak month / Trough month", "March (561) / November (437)", "Amplitude: 24% of mean"),
        ("Detrended SSN correlation", "r = 0.020, p = 0.753", "No residual solar signal"),
    ]
    for row_idx, (param, val, note) in enumerate(data, 1):
        table.rows[row_idx].cells[0].text = param
        table.rows[row_idx].cells[1].text = val
        table.rows[row_idx].cells[2].text = note
        for cell in table.rows[row_idx].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    return table


def create_manuscript():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # ══════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Brief Communication")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 102, 153)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        "Temporal Patterns in Pluripotent Stem Cell Research: "
        "An Exploratory Analysis of 6,101 GEO Datasets and Solar Activity Indices"
    )
    run.font.size = Pt(16)
    run.bold = True
    set_paragraph_format(p, space_after=Pt(18))

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Tatsuki Onishi")
    run.font.size = Pt(12)
    run2 = p.add_run("1,*")
    run2.font.superscript = True
    run2.font.size = Pt(9)
    set_paragraph_format(p, space_after=Pt(6))

    # Affiliations
    p = doc.add_paragraph()
    add_superscript_refs(p, "{1}[Affiliation to be added]")
    set_paragraph_format(p, space_after=Pt(3))

    p = doc.add_paragraph()
    run = p.add_run("*Correspondence: bougtoir@gmail.com")
    run.font.size = Pt(10)
    run.italic = True
    set_paragraph_format(p, space_after=Pt(18))

    # ══════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════
    add_heading(doc, "Abstract", level=2)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Background: Reproducibility in pluripotent stem cell (PSC) differentiation "
        "remains a persistent challenge, yet environmental variables beyond temperature "
        "and CO2 are rarely controlled or monitored. Here we explore whether temporal "
        "patterns exist in PSC research output that might reflect underlying seasonal "
        "or solar-geomagnetic influences. "
        "Methods: We queried the NCBI Gene Expression Omnibus (GEO) via Entrez E-utilities "
        "for all iPSC and ESC differentiation-related datasets (n = 6,101; 2001\u20132026). "
        "Metadata (submission date, institution country) were cross-referenced with NOAA "
        "solar cycle indices (F10.7 flux, sunspot number). Seasonal uniformity was tested "
        "by chi-square goodness-of-fit; solar correlations were assessed by Pearson "
        "correlation with 12-month rolling detrending. "
        "Results: Monthly dataset submissions deviated significantly from uniformity "
        "(\u03c7\u00b2 = 32.3, df = 11, p = 0.0007), peaking in March and troughing in November "
        "(amplitude 24% of mean). Geographic analysis revealed 97% Northern Hemisphere "
        "dominance, precluding hemisphere-inverted seasonal comparisons. Raw correlations "
        "with solar indices (r = 0.48 for F10.7) were entirely attributable to a secular "
        "growth trend; after detrending, the correlation vanished (r = 0.02, p = 0.75). "
        "Conclusions: GEO metadata confirm temporal non-uniformity in PSC research but "
        "cannot resolve whether this reflects biological seasonality or institutional "
        "cycles. The extreme hemisphere imbalance underscores the need for IVF registry "
        "data to test the photoperiod-driven hypothesis."
    )

    # Keywords
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(12))
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "pluripotent stem cells; seasonality; Gene Expression Omnibus; "
        "solar activity; reproducibility; environmental variables"
    )
    run.font.size = Pt(10)

    # ══════════════════════════════════════════════
    # MAIN TEXT
    # ══════════════════════════════════════════════
    doc.add_page_break()

    # --- Introduction ---
    add_heading(doc, "Introduction", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Directed differentiation of human pluripotent stem cells (PSCs) into "
        "specialized cell types holds transformative potential for regenerative "
        "medicine and disease modeling.{1} However, inter-laboratory reproducibility "
        "remains a well-documented challenge{2,3}: identical protocols applied in "
        "different facilities—or even the same facility at different times—often yield "
        "divergent outcomes. While genetic background, passage number, and reagent "
        "lot effects have been extensively studied,{13} environmental factors beyond "
        "the standard 37\u00b0C / 5% CO2 incubator setting receive remarkably little attention.{14}"
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Several lines of evidence suggest that seasonal environmental fluctuations "
        "may influence cell biology outcomes. In assisted reproduction, Australian data "
        "demonstrate that oocytes collected during summer yield ~30% higher live birth "
        "rates after frozen embryo transfer, with the effect tracking the hemisphere's "
        "photoperiod rather than the transfer season.{4} Murine embryos exhibit summer-specific "
        "two-cell block in the Northern Hemisphere.{5} At the molecular level, the circadian "
        "clock gene CRY1 directly regulates iPSC reprogramming efficiency,{6} and "
        "laboratory humidity—uncontrolled in most facilities—tracks seasonal cycles with "
        "documented impacts on contamination rates.{7}"
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "A separate hypothesis proposes that solar and geomagnetic activity—which "
        "follows ~11-year cycles and is synchronous across hemispheres—might influence "
        "cell proliferation.{8} Unlike photoperiod-driven effects (which should invert "
        "between hemispheres), geomagnetic effects would present as globally synchronous "
        "fluctuations."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "To provide an initial empirical assessment of these hypotheses, we conducted "
        "an exploratory meta-analysis of publicly available datasets in the NCBI Gene "
        "Expression Omnibus (GEO),{9} correlating PSC dataset submission patterns with "
        "seasonal cycles and NOAA solar activity indices.{10}"
    )

    # --- Methods ---
    add_heading(doc, "Methods", level=2)

    add_heading(doc, "Data sources", level=3)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "GEO metadata were retrieved via NCBI Entrez E-utilities (esearch.fcgi and "
        "esummary.fcgi) using the query: (iPSC OR \"induced pluripotent\" OR ESC OR "
        "\"embryonic stem cell\") AND (differentiation OR \"directed differentiation\" "
        "OR \"neural induction\" OR cardiomyocyte OR hepatocyte OR organoid). "
        "The search yielded 6,101 GEO Series records with associated submission dates, "
        "sample counts, organisms, and platform identifiers. "
        "For a stratified subsample (n = 600), contact country was extracted from GEO "
        "SOFT headers (Series_contact_country field) via direct HTTP requests to "
        "the GEO accession interface."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Solar activity indices (monthly F10.7 solar radio flux in solar flux units "
        "and international sunspot number) were obtained from the NOAA Space Weather "
        "Prediction Center (SWPC) observed solar cycle indices endpoint (JSON format).{10}"
    )

    add_heading(doc, "Statistical analysis", level=3)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Seasonal uniformity of monthly GEO submissions was tested using the "
        "chi-square goodness-of-fit test against the null hypothesis of equal "
        "submission probability across months. Countries were classified into "
        "Northern (\u226523.5\u00b0N) or Southern (\u226523.5\u00b0S) Hemisphere groups. "
        "Correlations between monthly dataset counts and solar indices were assessed "
        "using Pearson's r. To separate the secular growth trend in iPSC research "
        "from potential cyclic solar effects, a 12-month centered rolling mean was "
        "subtracted from both time series before computing the detrended correlation. "
        "All analyses were conducted in Python 3.11 (pandas, scipy.stats, matplotlib)."
    )

    # --- Results ---
    add_heading(doc, "Results", level=2)

    add_heading(doc, "Dataset characteristics", level=3)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The query returned 6,101 GEO Series spanning November 2001 to June 2026, "
        "representing a cumulative dataset of PSC differentiation experiments. "
        "Annual submissions grew from <10 in 2001\u20132005 to >600 per year after 2018, "
        "reflecting the exponential growth of iPSC research (Table 1)."
    )

    # Table 1
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(12), space_after=Pt(6))
    run = p.add_run("Table 1. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Summary of GEO iPSC/ESC differentiation dataset characteristics "
                    "and key statistical results.")
    run.font.size = Pt(10)

    add_table(doc)

    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(12))

    add_heading(doc, "Seasonal distribution", level=3)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Monthly distribution of GEO submissions deviated significantly from "
        "uniformity (\u03c7\u00b2 = 32.3, df = 11, p = 0.0007; Figure 1A). The peak "
        "month was March (n = 561, 9.2%) and the trough was November (n = 437, 7.2%), "
        "yielding an amplitude of 124 datasets (24% of the monthly mean of 508.4). "
        "A secondary peak in December\u2013January was also observed. This pattern is "
        "broadly consistent with Northern Hemisphere academic cycles (grant deadlines, "
        "publication timing), but cannot distinguish institutional from biological "
        "seasonality without additional experimental metadata."
    )

    add_heading(doc, "Geographic distribution and hemisphere balance", level=3)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Among the 600 datasets with retrievable country data, 584 (97.3%) originated "
        "from Northern Hemisphere institutions and only 16 (2.7%) from Southern Hemisphere "
        "countries (Figure 1B). The top five contributing countries were the United States "
        "(n = 265, 44%), China (n = 59, 10%), United Kingdom (n = 44, 7%), Germany "
        "(n = 43, 7%), and Japan (n = 36, 6%). Southern Hemisphere representation was "
        "limited to Australia (n = 13), Brazil (n = 2), and South Africa (n = 1). "
        "This extreme hemispheric imbalance precludes a meaningful test of the "
        "hemisphere-inverted seasonal hypothesis using GEO data alone."
    )

    # Insert Figure 1
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(18), space_after=Pt(6))

    fig1_path = os.path.join(OUTPUT_DIR, "geo_solar_comprehensive.png")
    if os.path.exists(fig1_path):
        doc.add_picture(fig1_path, width=Inches(6.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(12), space_after=Pt(12))
    run = p.add_run("Figure 1. ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        "Comprehensive analysis of temporal and geographic patterns in GEO PSC "
        "differentiation datasets. (A) Monthly distribution with chi-square test "
        "for uniformity. (B) Geographic distribution of datasets by country. "
        "(C) Normalized monthly distribution by hemisphere. (D) Annual dataset "
        "trend overlaid with solar cycle (sunspot number). (E) Monthly dataset "
        "count vs. F10.7 solar radio flux (raw correlation). (F) Detrended "
        "correlation between dataset count and sunspot number residuals."
    )
    run.font.size = Pt(10)

    add_heading(doc, "Solar activity correlations", level=3)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Raw Pearson correlation between monthly GEO dataset counts and solar "
        "activity indices was positive and highly significant: r = 0.478 (p < 10\u207b\u00b9\u2074) "
        "for F10.7 flux and r = 0.402 (p < 10\u207b\u00b9\u2070) for sunspot number "
        "(Figure 1E). However, visual inspection of the time series (Figure 1D) revealed "
        "that both iPSC research output and solar activity increased during the same "
        "2010\u20132025 window, suggesting confounding by a shared secular trend."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "After subtracting the 12-month rolling mean from both series, the detrended "
        "correlation was negligible (r = 0.020, p = 0.753; Figure 1F), confirming that "
        "the raw association was spurious. No residual relationship between solar "
        "activity fluctuations and PSC dataset submission timing was detectable at "
        "this level of analysis."
    )

    # --- Discussion ---
    add_heading(doc, "Discussion", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "This exploratory analysis of 6,101 GEO datasets yields three principal findings. "
        "First, PSC differentiation dataset submissions are non-uniformly distributed "
        "across the calendar year, with a March peak and November trough. Second, the "
        "global PSC research landscape is overwhelmingly concentrated in the Northern "
        "Hemisphere (~97%), rendering GEO metadata insufficient for testing hemisphere-"
        "inverted seasonal hypotheses. Third, the apparent correlation between research "
        "output and solar activity is entirely attributable to coincident secular trends, "
        "providing no evidence for a direct geomagnetic influence on dataset submissions."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "Several important limitations constrain interpretation. Most critically, "
        "GEO submission date is a proxy for—not a measure of—experiment execution "
        "date. The typical delay between experiment completion and data deposition "
        "(estimated 6\u201318 months) introduces substantial temporal noise that would "
        "obscure genuine seasonal effects on cell biology. Additionally, the observed "
        "monthly pattern likely reflects institutional rhythms (academic calendar, "
        "grant cycles, conference submission deadlines) rather than biological seasonality "
        "per se. Disentangling these requires access to actual experiment dates, which "
        "are not captured in GEO metadata."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The hemisphere imbalance has important implications for study design. "
        "The photoperiod-driven hypothesis{4,6,12} predicts that seasonal effects on "
        "cell differentiation should invert between hemispheres (i.e., summer advantage "
        "in June\u2013August for Northern Hemisphere and December\u2013February for Southern "
        "Hemisphere). Testing this \"natural experiment\" requires substantial Southern "
        "Hemisphere data, which GEO cannot provide. IVF registries—particularly the "
        "HFEA (United Kingdom; >1 million cycles) and ANZARD (Australia/New Zealand; "
        "~109,000 cycles/year)—offer cycle-level data with precise treatment dates "
        "in both hemispheres and represent the most promising resource for this "
        "hypothesis test."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "The null result for the solar/geomagnetic hypothesis does not rule out "
        "such effects on cell biology{8}—it merely demonstrates that GEO submission "
        "metadata are too noisy and temporally displaced to detect them. Direct "
        "measurements of geomagnetic disturbance at the time of cell culture "
        "experiments{11} would be required for a definitive test."
    )

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    add_superscript_refs(p,
        "In conclusion, this analysis provides suggestive evidence of temporal "
        "structure in PSC research output and highlights the fundamental limitations "
        "of publicly available metadata for testing environmental hypotheses. We "
        "propose that systematic environmental profiling of stem cell facilities{14}—"
        "including continuous monitoring of humidity, ambient light, electromagnetic "
        "fields, and atmospheric pressure—combined with retrospective analysis of "
        "IVF registry data represents the most viable path toward identifying and "
        "controlling the \"invisible variables\" of stem cell culture.{15}"
    )

    # ══════════════════════════════════════════════
    # DATA AVAILABILITY
    # ══════════════════════════════════════════════
    add_heading(doc, "Data Availability", level=2)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=1.5)
    run = p.add_run(
        "All GEO metadata and NOAA solar indices used in this study are publicly "
        "available. Analysis scripts and processed data are deposited at "
        "https://github.com/bougtoir/wip (stem_cell_seasonality/ directory)."
    )

    # ══════════════════════════════════════════════
    # ACKNOWLEDGMENTS
    # ══════════════════════════════════════════════
    add_heading(doc, "Acknowledgments", level=2)
    p = doc.add_paragraph()
    run = p.add_run("[To be added]")
    run.italic = True
    set_paragraph_format(p, line_spacing=1.5)

    # ══════════════════════════════════════════════
    # DECLARATION OF INTERESTS
    # ══════════════════════════════════════════════
    add_heading(doc, "Declaration of Interests", level=2)
    p = doc.add_paragraph()
    run = p.add_run("The author declares no competing interests.")
    set_paragraph_format(p, line_spacing=1.5)

    # ══════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "References", level=2)

    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=Pt(3), line_spacing=1.15)
        run = p.add_run(f"{i}. ")
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(ref)
        run.font.size = Pt(9)

    # ══════════════════════════════════════════════
    # Save
    # ══════════════════════════════════════════════
    out_path = os.path.join(OUTPUT_DIR, "GEO_Solar_Analysis_BriefComm.docx")
    doc.save(out_path)
    print(f"Manuscript saved to: {out_path}")

    # Word count estimate
    word_count = 0
    for para in doc.paragraphs:
        word_count += len(para.text.split())
    print(f"Approximate word count (all text): {word_count}")

    return out_path


if __name__ == "__main__":
    create_manuscript()
