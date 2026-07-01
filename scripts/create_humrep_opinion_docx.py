#!/usr/bin/env python3
"""
Generate Human Reproduction Opinion manuscript:
"The Invisible Variables: Uncontrolled Environmental Factors in
 Stem Cell Differentiation and the Fiscal-Year Artifact in Public Databases"

Format: Human Reproduction Opinion
- Author-date citation style (Author et al., Year) in text
- References listed alphabetically
- Up to 10 authors before et al.
- Unstructured abstract (≤300 words)
- 5–10 keywords after abstract
- Graphical abstract (separate file)
- Declaration of authors' roles after Acknowledgements
- Funding statement
- Conflict of interest statement
- Figures uploaded separately; figure legends at end of manuscript
- Double-spaced, Times New Roman 12pt

Reference format (Paperpile / Human Reproduction style):
  Author1 A, Author2 B, Author3 C et al. Title. Journal Year;Vol:Pages.
In-text: (Author, Year) or (Author1 and Author2, Year) or (Author et al., Year)
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ──────────────────────────────────────────────
# Reference database
# Format: dict with key = cite_key, fields for author-date style
# All references verified against PubMed/DOI 2026-06-10
# ──────────────────────────────────────────────
REFERENCES = [
    {
        "key": "Kirkeby2025",
        "authors": "Kirkeby A, Main H, Carpenter M",
        "year": 2025,
        "title": "Pluripotent stem-cell-derived therapies in clinical trial: a 2025 update",
        "journal": "Cell Stem Cell",
        "volume": "32",
        "issue": "1",
        "pages": "10–37",
        "cite": "Kirkeby et al., 2025",
    },
    {
        "key": "Yamanaka2020",
        "authors": "Yamanaka S",
        "year": 2020,
        "title": "Pluripotent stem cell-based cell therapy—promise and challenges",
        "journal": "Cell Stem Cell",
        "volume": "27",
        "issue": "4",
        "pages": "523–531",
        "cite": "Yamanaka, 2020",
    },
    {
        "key": "Volpato2018",
        "authors": "Volpato V, Smith J, Sandor C et al.",
        "year": 2018,
        "title": "Reproducibility of molecular phenotypes after long-term "
                 "differentiation to human iPSC-derived neurons: a multi-site omics study",
        "journal": "Stem Cell Reports",
        "volume": "11",
        "issue": "4",
        "pages": "897–911",
        "cite": "Volpato et al., 2018",
    },
    {
        "key": "VolpatoWebber2020",
        "authors": "Volpato V, Webber C",
        "year": 2020,
        "title": "Addressing variability in iPSC-derived models of human disease: "
                 "guidelines to promote reproducibility",
        "journal": "Dis Model Mech",
        "volume": "13",
        "issue": "1",
        "pages": "dmm042317",
        "cite": "Volpato and Webber, 2020",
    },
    {
        "key": "Ortmann2017",
        "authors": "Ortmann D, Vallier L",
        "year": 2017,
        "title": "Variability of human pluripotent stem cell lines",
        "journal": "Curr Opin Genet Dev",
        "volume": "46",
        "pages": "179–185",
        "cite": "Ortmann and Vallier, 2017",
    },
    {
        "key": "Cai2025",
        "authors": "Cai J, Li X, Zhang Y et al.",
        "year": 2025,
        "title": "Real-time monitoring reveals the effects of low concentrations "
                 "of volatile organic compounds in the embryology laboratory",
        "journal": "Hum Reprod",
        "volume": "40",
        "issue": "4",
        "pages": "deaf008",
        "cite": "Cai et al., 2025",
    },
    {
        "key": "Agarwal2017",
        "authors": "Agarwal N, Chattopadhyay R, Ghosh S et al.",
        "year": 2017,
        "title": "Volatile organic compounds and good laboratory practices in the "
                 "in vitro fertilization laboratory: the important parameters for "
                 "successful outcome in extended culture",
        "journal": "J Assist Reprod Genet",
        "volume": "34",
        "issue": "8",
        "pages": "999–1006",
        "cite": "Agarwal et al., 2017",
    },
    {
        "key": "Panina2021",
        "authors": "Panina Y, Yamane J, Kobayashi K, Sone H, Fujibuchi W",
        "year": 2021,
        "title": "Human ES and iPS cells display less drug resistance than "
                 "differentiated cells, and naive-state induction further decreases "
                 "drug resistance",
        "journal": "J Toxicol Sci",
        "volume": "46",
        "issue": "3",
        "pages": "131–142",
        "cite": "Panina et al., 2021",
    },
    {
        "key": "McCreery2025",
        "authors": "McCreery KP, Stubb A, Stephens R et al.",
        "year": 2025,
        "title": "Mechano-osmotic signals control chromatin state and fate "
                 "transitions in pluripotent stem cells",
        "journal": "Nat Cell Biol",
        "volume": "27",
        "issue": "10",
        "pages": "1757–1770",
        "cite": "McCreery et al., 2025",
    },
    {
        "key": "Sato2023",
        "authors": "Sato S, Hishida T, Kinouchi K et al.",
        "year": 2023,
        "title": "The circadian clock CRY1 regulates pluripotent stem cell "
                 "identity and somatic cell reprogramming",
        "journal": "Cell Rep",
        "volume": "42",
        "issue": "6",
        "pages": "112590",
        "cite": "Sato et al., 2023",
    },
    {
        "key": "Ameneiro2020",
        "authors": "Ameneiro C, Moreira T, Fuentes-Iglesias A et al.",
        "year": 2020,
        "title": "BMAL1 coordinates energy metabolism and differentiation of "
                 "pluripotent stem cells",
        "journal": "Life Sci Alliance",
        "volume": "3",
        "issue": "5",
        "pages": "e201900534",
        "cite": "Ameneiro et al., 2020",
    },
    {
        "key": "Golan2019",
        "authors": "Golan K, Kollet O, Markus RP, Lapidot T",
        "year": 2019,
        "title": "Daily light and darkness onset and circadian rhythms metabolically "
                 "synchronize hematopoietic stem cell differentiation and maintenance",
        "journal": "Exp Hematol",
        "volume": "78",
        "pages": "1–10",
        "cite": "Golan et al., 2019",
    },
    {
        "key": "Bi2020",
        "authors": "Bi S, Tang J, Zhang L et al.",
        "year": 2020,
        "title": "Fine particulate matter reduces the pluripotency and proliferation "
                 "of human embryonic stem cells through ROS induced AKT and ERK "
                 "signaling pathway",
        "journal": "Reprod Toxicol",
        "volume": "96",
        "pages": "231–240",
        "cite": "Bi et al., 2020",
    },
    {
        "key": "Chui2024",
        "authors": "Chui JS-H, Izuel-Idoype T, Qualizza A et al.",
        "year": 2024,
        "title": "Osmolar modulation drives reversible cell cycle exit and human "
                 "pluripotent cell differentiation via NF-\u03baB and WNT signaling",
        "journal": "Adv Sci",
        "volume": "11",
        "issue": "7",
        "pages": "2307554",
        "cite": "Chui et al., 2024",
    },
    {
        "key": "Mizuno2020",
        "authors": "Mizuno M, Endo K, Katano H et al.",
        "year": 2020,
        "title": "The environmental risk assessment of cell-processing facilities "
                 "for cell therapy in a Japanese academic institution",
        "journal": "PLoS One",
        "volume": "15",
        "issue": "8",
        "pages": "e0236600",
        "cite": "Mizuno et al., 2020",
    },
    {
        "key": "Klein2022",
        "authors": "Klein SG, Steckbauer A, Alsolami SM et al.",
        "year": 2022,
        "title": "Toward best practices for controlling mammalian cell culture environments",
        "journal": "Front Cell Dev Biol",
        "volume": "10",
        "pages": "788808",
        "cite": "Klein et al., 2022",
    },
    {
        "key": "Czyz2004",
        "authors": "Czyz J, Nikolova T, Schuderer J et al.",
        "year": 2004,
        "title": "Non-thermal effects of power-line magnetic fields (50 Hz) on "
                 "gene expression levels of pluripotent embryonic stem cells",
        "journal": "Mutat Res",
        "volume": "557",
        "issue": "1",
        "pages": "63–74",
        "cite": "Czyz et al., 2004",
    },
    {
        "key": "Diatroptova2022",
        "authors": "Diatroptova MA, Kosyreva AM, Diatroptov ME",
        "year": 2022,
        "title": "About 4-day rhythm of proliferative activity of L-929 cells "
                 "in culture correlates with the intensity of secondary cosmic "
                 "radiation fluctuations",
        "journal": "Bull Exp Biol Med",
        "volume": "172",
        "issue": "5",
        "pages": "561–565",
        "cite": "Diatroptova et al., 2022",
    },
    {
        "key": "Leathersich2023",
        "authors": "Leathersich SJ, Hart RJ, Wijs LA et al.",
        "year": 2023,
        "title": "Season at the time of oocyte collection and frozen embryo "
                 "transfer outcomes",
        "journal": "Hum Reprod",
        "volume": "38",
        "issue": "9",
        "pages": "1714–1722",
        "cite": "Leathersich et al., 2023",
    },
    {
        "key": "Wang2025",
        "authors": "Wang C, Chen J, Lin Z et al.",
        "year": 2025,
        "title": "The impact of season, temperature, and direct normal irradiance "
                 "on IVF pregnancy outcomes: a retrospective cohort study",
        "journal": "Int J Biometeorol",
        "volume": "69",
        "pages": "2053–2065",
        "cite": "Wang et al., 2025",
    },
    {
        "key": "Braga2025",
        "authors": "Braga DPA, Setti A, Guilherme P, Iaconelli A Jr, Borges E Jr",
        "year": 2025,
        "title": "Association between meteorological season and embryo quality "
                 "in the era of morphokinetics",
        "journal": "J Assist Reprod Genet",
        "volume": "42",
        "pages": "1287–1298",
        "cite": "Braga et al., 2025",
    },
    {
        "key": "Deng2025",
        "authors": "Deng Q, Wu F, Wang J et al.",
        "year": 2025,
        "title": "Association between season and pregnancy outcomes in fresh "
                 "embryo transfer cycles: a systematic review and meta-analysis",
        "journal": "Front Public Health",
        "volume": "13",
        "pages": "1660982",
        "cite": "Deng et al., 2025",
    },
    {
        "key": "Wang1992",
        "authors": "Wang JP, Her WY, Meir YJ, Liu TS, Chang HL, Haung FL",
        "year": 1992,
        "title": "Seasonal variation in cell cycle during early development "
                 "of the mouse embryo",
        "journal": "Reproduction",
        "volume": "94",
        "issue": "2",
        "pages": "431–436",
        "cite": "Wang et al., 1992",
    },
    {
        "key": "Barrett2013",
        "authors": "Barrett T, Wilhite SE, Ledoux P et al.",
        "year": 2013,
        "title": "NCBI GEO: archive for functional genomics data sets—update",
        "journal": "Nucleic Acids Res",
        "volume": "41",
        "issue": "D",
        "pages": "D991–D995",
        "cite": "Barrett et al., 2013",
    },
    {
        "key": "MardiaJupp2000",
        "authors": "Mardia KV, Jupp PE",
        "year": 2000,
        "title": "Directional Statistics",
        "journal": "",
        "volume": "",
        "pages": "",
        "publisher": "Wiley",
        "type": "book",
        "cite": "Mardia and Jupp, 2000",
    },
    {
        "key": "Karagiannis2019",
        "authors": "Karagiannis P, Takahashi K, Saito M et al.",
        "year": 2019,
        "title": "Induced pluripotent stem cells and their use in human models "
                 "of disease and development",
        "journal": "Physiol Rev",
        "volume": "99",
        "issue": "1",
        "pages": "79–114",
        "cite": "Karagiannis et al., 2019",
    },
]


def format_ref_for_list(ref):
    """Format a reference for the alphabetical reference list (Human Reproduction style)."""
    if ref.get("type") == "book":
        return f"{ref['authors']}. {ref['title']}. {ref['year']}; {ref['publisher']}."

    parts = ref["authors"]
    parts += f". {ref['title']}. {ref['journal']} {ref['year']}"
    if ref.get("volume"):
        parts += f";{ref['volume']}"
    if ref.get("pages"):
        parts += f":{ref['pages']}"
    parts += "."
    return parts


def get_sorted_references():
    """Return references sorted alphabetically by first author surname."""
    def sort_key(ref):
        surname = ref["authors"].split(",")[0].split()[-1] if "," in ref["authors"] else ref["authors"].split()[0]
        return (surname.lower(), ref["year"])
    return sorted(REFERENCES, key=sort_key)


def set_paragraph_format(para, space_after=Pt(6), space_before=Pt(0),
                         line_spacing=2.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    para.paragraph_format.line_spacing = line_spacing
    para.alignment = alignment
    return para


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_body_text(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p


def create_manuscript():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ──────────────────────────────────────────────
    # Title page
    # ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Opinion")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 102, 153)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(24), space_after=Pt(12))
    run = p.add_run(
        "The Invisible Variables: Uncontrolled Environmental Factors in "
        "Stem Cell Differentiation and the Fiscal-Year Artifact in Public Databases"
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(24))
    run = p.add_run("[Author names and affiliations to be added]")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(6))
    run = p.add_run("Correspondence: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run("[Corresponding author name, full postal address, email, ORCID iD]")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True

    # ──────────────────────────────────────────────
    # Abstract
    # ──────────────────────────────────────────────
    add_heading(doc, "Abstract", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run(
        "Despite stringent control of temperature and CO\u2082, pluripotent stem cell (PSC) "
        "differentiation remains plagued by unexplained batch-to-batch variability. We "
        "argue that multiple environmental variables\u2014humidity, volatile organic compounds "
        "(VOCs), ambient light, and electromagnetic fields (EMF)\u2014remain unmonitored in "
        "most laboratories and may constitute controllable sources of variation. Evidence "
        "shows that PSCs are uniquely vulnerable to such perturbations: iPSCs display "
        "several-fold greater chemical sensitivity than somatic cells, and osmotic stress "
        "suffices to alter chromatin state and accelerate differentiation. The IVF field "
        "provides proof-of-concept: oocytes collected in summer yield 30% higher live birth "
        "rates than those collected in autumn, and VOC reduction increased blastocyst rates "
        "by 18%. To test whether seasonal variation leaves a footprint in research output, "
        "we analyzed 6,101 GEO PSC datasets. A natural experiment exploiting geographic "
        "variation in academic year calendars reveals that the apparent March peak is an "
        "institutional artifact of Japan's fiscal year-end, not a biological signal. We "
        "propose that prospective multi-parameter environmental monitoring\u2014not database "
        "mining\u2014is the path to identifying which invisible variables matter."
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # ──────────────────────────────────────────────
    # Keywords
    # ──────────────────────────────────────────────
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(12), space_after=Pt(12))
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run(
        "pluripotent stem cells; differentiation variability; environmental monitoring; "
        "seasonality; volatile organic compounds; IVF laboratory; GEO database; "
        "reproducibility"
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # ──────────────────────────────────────────────
    # Section 1: Introduction
    # ──────────────────────────────────────────────
    add_heading(doc, "The Reproducibility Crisis Has an Environmental Blind Spot", level=2)

    add_body_text(doc,
        "Pluripotent stem cell (PSC)-derived therapies are advancing rapidly toward "
        "clinical application, with over 115 trials and 1,200 patients dosed as of "
        "2024 (Kirkeby et al., 2025). Yet differentiation protocols remain notoriously "
        "variable across laboratories and even between batches within a single group "
        "(Yamanaka, 2020; Volpato et al., 2018; Volpato and Webber, 2020; "
        "Ortmann and Vallier, 2017). The field has invested enormously in optimizing "
        "controllable recipe variables\u2014growth factors, small molecules, extracellular "
        "matrices, and timing\u2014while largely ignoring the kitchen in which these recipes "
        "are executed."
    )

    add_body_text(doc,
        "Consider an analogy from artisanal baking. For decades, bread quality varied "
        "unpredictably between batches until industrial bakeries recognized that controlling "
        "only oven temperature was insufficient; ambient humidity, dough temperature during "
        "proofing, and air quality were equally critical. Today's PSC laboratory resembles "
        "the pre-industrial bakery: we control incubator temperature (37\u00b0C) and CO\u2082 (5%) "
        "but leave humidity, volatile organic compounds (VOCs), ambient light, electromagnetic "
        "fields (EMF), and barometric pressure entirely unmonitored. This Opinion argues "
        "that these 'invisible variables' may constitute a significant and correctable source "
        "of differentiation variability."
    )

    add_body_text(doc,
        "A deeper methodological issue compounds the problem. Research using clonal systems\u2014"
        "isogenic cell lines, inbred mice, genetically identical iPSC clones\u2014operates under "
        "the implicit assumption that genetic homogeneity renders complex confounding "
        "adjustment unnecessary. Because the genetic background is 'controlled,' "
        "investigators rarely apply the stratification, regression, or sensitivity analyses "
        "that epidemiologists would consider mandatory for observational data. This creates "
        "a blind spot: non-genetic confounders (environmental, institutional, temporal) are "
        "neither measured nor adjusted for, because the clonal paradigm implies they should "
        "not exist. Our analysis reveals that they do exist\u2014and that aggregate data from "
        "clonal systems can be just as non-homogeneous and confounded as observational "
        "cohort data when institutional and environmental variation is ignored."
    )

    add_body_text(doc,
        "The scale of the problem is not trivial. Multi-site reproducibility studies "
        "consistently report substantial inter-laboratory variation in PSC differentiation "
        "outcomes even when nominally identical protocols are followed (Volpato et al., 2018). "
        "Volpato et al. found that laboratory-of-origin was the dominant source of variance "
        "in iPSC-derived neuron transcriptomes, exceeding the effect of genetic background or "
        "differentiation batch. Current explanations invoke differences in operator technique, "
        "reagent lots, and passage number\u2014but these account for only a fraction of the "
        "observed variance. We propose that uncontrolled environmental differences between "
        "laboratories constitute a missing explanatory variable."
    )

    add_body_text(doc,
        "This proposal is neither speculative nor unprecedented. The in vitro fertilization "
        "(IVF) field underwent a parallel revolution when it recognized that air quality in "
        "embryology laboratories was a major determinant of blastocyst formation rates "
        "(Cai et al., 2025; Agarwal et al., 2017). Prior to this insight, unexplained "
        "clinic-to-clinic variation in IVF success rates was attributed to operator skill "
        "and patient selection; after implementing VOC filtration and environmental controls, "
        "many laboratories saw 10\u201320% improvements in outcomes (Agarwal et al., 2017). We "
        "argue that the PSC field is at a similar inflection point, with the added advantage "
        "that modern IoT sensor technology makes comprehensive environmental monitoring far "
        "cheaper and easier than it was for IVF laboratories two decades ago."
    )

    # ──────────────────────────────────────────────
    # Section 2: Why PSCs Are Uniquely Vulnerable
    # ──────────────────────────────────────────────
    add_heading(doc, "Why Pluripotent Stem Cells Are Uniquely Vulnerable", level=2)

    add_body_text(doc,
        "The premise that uncontrolled environmental factors could meaningfully affect PSC "
        "differentiation\u2014while being negligible for established cell lines\u2014requires "
        "justification. Several lines of evidence support this claim."
    )

    add_body_text(doc,
        "First, PSCs display markedly greater chemical sensitivity than differentiated "
        "cells. Panina et al. (2021) demonstrated that human iPS cells are approximately "
        "1.5-fold more sensitive to drug exposure than ES cells, and both are several-fold "
        "less resistant than non-pluripotent cell types. Critically, naive-state "
        "induction\u2014pushing cells toward a more primitive pluripotent state\u2014further increased "
        "sensitivity, establishing a gradient: the more undifferentiated the cell, the more "
        "vulnerable it is to environmental perturbation."
    )

    add_body_text(doc,
        "Second, osmotic perturbations of surprisingly small magnitude can gate cell fate "
        "transitions. McCreery et al. (2025) showed that compaction-triggered changes in "
        "nuclear shape and volume in human iPSCs remodel chromatin architecture, increase "
        "nucleoplasmic viscosity, and prime cells for ectodermal differentiation. They "
        "propose that 'mechano-osmotic reprogramming of the nuclear environment tunes "
        "differentiation efficiency by lowering the energy barrier for cell fate "
        "transitions.' This finding has immediate implications for humidity control: during "
        "biosafety cabinet work, culture medium evaporates at rates determined by ambient "
        "humidity, producing osmolarity increases that could cross the thresholds identified "
        "in their work."
    )

    add_body_text(doc,
        "Third, PSCs depend on circadian clock components in a non-canonical manner. CRY1 is "
        "dramatically upregulated in iPSCs and ESCs compared to somatic cells, and its "
        "deletion impairs self-renewal and disrupts differentiation (Sato et al., 2023). "
        "BMAL1 coordinates energy metabolism and differentiation independently of circadian "
        "oscillation, which is suppressed in pluripotent cells (Ameneiro et al., 2020). "
        "Consistent with this, daily light and darkness onset metabolically synchronize "
        "hematopoietic stem cell (HSC) differentiation and maintenance (Golan et al., 2019), "
        "suggesting that light-sensitive circadian pathways are broadly relevant to stem cell "
        "biology beyond the pluripotent compartment. This architecture\u2014dependence on clock "
        "molecules without functional clock oscillation\u2014may render PSCs vulnerable to "
        "external zeitgeber inputs (e.g., ambient light leaking into incubators or biosafety "
        "cabinets) that would not perturb differentiated cells with intact circadian buffering."
    )

    add_body_text(doc,
        "Fourth, airborne pollutants affect PSCs at concentrations that leave somatic cells "
        "unharmed. Fine particulate matter (PM2.5) downregulates pluripotency markers "
        "(NANOG, OCT4) in human ESCs through ROS-mediated AKT/ERK signaling (Bi et al., "
        "2020). In IVF laboratories, real-time VOC monitoring revealed that even low "
        "concentrations linearly predicted decreased blastocyst quality, and reducing VOCs "
        "increased blastocyst formation by 18% and live birth rates by 8% (Cai et al., 2025; "
        "Agarwal et al., 2017). Importantly, the effect was specific to fresh embryos\u2014"
        "frozen embryos were unaffected\u2014indicating that early-stage cells are selectively "
        "vulnerable."
    )

    add_body_text(doc,
        "Fifth, osmolarity modulation through hyperosmotic culture drives NF-\u03baB and WNT "
        "signaling-dependent cell cycle exit and maturation in iPSC-derived cells "
        "(Chui et al., 2024). This means that osmolarity changes of physiological "
        "magnitude\u2014such as those caused by medium evaporation during cabinet work\u2014are not "
        "merely stress events but active differentiation signals that can override intended "
        "protocol cues. The key point is not that any single perturbation is catastrophic, "
        "but that PSCs sit at an energetically shallow decision landscape where subtle "
        "environmental shifts can tip the balance between self-renewal and lineage commitment."
    )

    add_body_text(doc,
        "Together, these findings establish a coherent picture: PSCs are intrinsically more "
        "sensitive to environmental perturbation than differentiated cells because of their "
        "open chromatin state, rapid proliferation, dependence on clock-associated proteins, "
        "and susceptibility to osmotic gating of fate decisions. What is unknown is the "
        "real-world magnitude of these effects in working laboratories. This gap exists not "
        "because the question is unanswerable, but because no one has yet systematically "
        "measured the relevant variables alongside differentiation outcomes."
    )

    # ──────────────────────────────────────────────
    # Section 3: Catalogue of Uncontrolled Variables
    # ──────────────────────────────────────────────
    add_heading(doc, "A Catalogue of Uncontrolled Environmental Variables", level=2)

    add_body_text(doc,
        "A comprehensive environmental assessment of a Japanese cell-processing facility "
        "revealed that while temperature was maintained constant year-round, humidity "
        "tracked outdoor seasonal patterns, peaking in summer and reaching troughs in "
        "winter (Mizuno et al., 2020). Bacterial and fungal contamination rates were "
        "significantly elevated at humidity above 55%. The authors noted that humidity "
        "control equipment 'is expensive and usually not set up in academic institutions.' "
        "Industry guidelines confirm that seasonal transitions substantially increase "
        "microbial contamination risk (Klein et al., 2022)."
    )

    add_body_text(doc,
        "VOC concentrations in laboratories are determined by building materials, "
        "disinfectants, and HVAC systems, with off-gassing rates that are strongly "
        "temperature- and humidity-dependent\u2014introducing a seasonal component (Cai et al., "
        "2025; Agarwal et al., 2017). Ambient light exposure occurs during every biosafety "
        "cabinet manipulation, with duration and intensity varying by laboratory design and "
        "season (longer daylight hours in summer). Given CRY1's role in maintaining "
        "pluripotency (Sato et al., 2023), even brief light exposure during passage or "
        "feeding could introduce stochastic variation in differentiation priming. "
        "Electromagnetic fields from incubator heaters, freezer compressors, and building "
        "wiring (50/60 Hz) have been shown to alter gene expression in mouse ESCs "
        "(Czyz et al., 2004), though effect sizes under typical laboratory conditions "
        "remain uncharacterized."
    )

    add_body_text(doc,
        "Each of these variables exhibits characteristic seasonal patterns. Humidity tracks "
        "outdoor climate in non-climate-controlled clean rooms (Mizuno et al., 2020). VOC "
        "off-gassing rates increase with temperature and are modulated by HVAC cycling "
        "patterns that change between heating and cooling seasons. Photoperiod varies by "
        "latitude, with laboratories at high latitudes experiencing >8 hours of daylight "
        "difference between solstices. Barometric pressure fluctuations correlate with storm "
        "frequency, which is itself seasonal. Cosmic ray flux shows an 11-year solar cycle "
        "modulation with additional seasonal components (Diatroptova et al., 2022). The "
        "combination of these seasonal environmental signals creates a complex, "
        "latitude-dependent forcing function that no PSC laboratory currently monitors."
    )

    add_body_text(doc,
        "Barometric pressure fluctuations deserve particular attention. Low-pressure "
        "weather systems reduce the partial pressure of dissolved gases in culture media "
        "and alter gas exchange kinetics in non-equilibrated incubators. These fluctuations "
        "are acute (hours), large (up to 30 hPa during cyclones), and entirely uncompensated "
        "in standard CO\u2082 incubators. Although direct evidence linking barometric pressure to "
        "PSC differentiation is lacking, the sensitivity of pH-dependent signaling pathways "
        "(e.g., Wnt, Hedgehog) to dissolved CO\u2082 concentrations suggests a plausible "
        "mechanism. Similarly, building vibration from HVAC systems, foot traffic, and "
        "nearby construction transmits mechanical stress to cultured cells\u2014and recent work "
        "demonstrates that mechanical compression can modulate nuclear mechanics and "
        "differentiation timing in hiPSCs (McCreery et al., 2025)."
    )

    # ──────────────────────────────────────────────
    # Section 4: Seasonality in IVF (external evidence)
    # ──────────────────────────────────────────────
    add_heading(doc, "Seasonality Evidence from IVF: A Proof of Concept", level=2)

    add_body_text(doc,
        "The IVF field provides the strongest existing evidence that season affects "
        "early-stage cell outcomes. Leathersich et al. (2023) analyzed 3,659 frozen embryo "
        "transfer cycles in Perth, Australia (33\u00b0S), finding that oocytes collected in "
        "summer had 30% higher live birth rates than those collected in autumn. Crucially, "
        "the season of embryo transfer was irrelevant\u2014only the season at oocyte collection "
        "mattered\u2014establishing that the effect acts on the gamete/early embryo rather than "
        "on implantation. Northern Hemisphere studies report concordant summer "
        "advantages (Wang et al., 2025), and a Brazilian study at 23\u00b0S confirms the "
        "expected six-month phase shift in embryo quality metrics (Braga et al., 2025). "
        "However, a recent meta-analysis of 159,696 fresh transfers found minimal overall "
        "effects (Deng et al., 2025), suggesting the impact is specific to oocyte/embryo "
        "quality rather than implantation."
    )

    add_body_text(doc,
        "Wang et al. (1992) provided direct experimental evidence in mice: embryos cultured "
        "during summer showed the 'two-cell block' phenomenon at the early two-cell stage, "
        "while winter embryos developed normally. Cleavage from two-cell to four-cell was "
        "also delayed in summer. This finding\u2014a seasonal effect on preimplantation embryo "
        "development under ostensibly controlled culture conditions\u2014is precisely the type of "
        "signal that uncontrolled environmental variables would produce."
    )

    add_body_text(doc,
        "The IVF evidence is particularly instructive because it provides a causal "
        "structure: the season at oocyte collection affects outcomes, but the season at "
        "embryo transfer does not (Leathersich et al., 2023). This temporal dissection "
        "implies that the environmental effect acts during a specific developmental "
        "window\u2014oocyte maturation and early cleavage\u2014rather than through a general mechanism "
        "affecting all cells at all times. PSC differentiation protocols, which recapitulate "
        "aspects of early embryonic lineage commitment over 7\u201330 days, would be expected to "
        "share this window-specific vulnerability. The parallel is not merely analogical: "
        "iPSCs undergoing directed differentiation traverse epigenetic states homologous to "
        "those of the early embryo, including the exit from naive pluripotency, lineage "
        "priming, and specification\u2014each potentially gated by the same environmental "
        "sensitivities that affect preimplantation development."
    )

    # ──────────────────────────────────────────────
    # Section 5: The GEO Natural Experiment
    # ──────────────────────────────────────────────
    add_heading(doc, "A Natural Experiment Reveals Institutional, Not Biological, Seasonality",
                level=2)

    add_body_text(doc,
        "Given the evidence above, we asked whether seasonal patterns are detectable in "
        "public PSC research databases. We analyzed submission dates for 6,101 PSC "
        "differentiation datasets deposited in NCBI GEO (Barrett et al., 2013) between "
        "2001 and 2026 (search query: 'iPSC OR ESC differentiation'; inclusion: all dataset "
        "types with valid submission dates; exclusion: datasets lacking month-level date "
        "resolution). The overall distribution was significantly non-uniform "
        "(\u03c7\u00b2=32.3, df=11, p=0.0007) with a pronounced March peak\u2014superficially consistent "
        "with a spring photoperiod advantage in the Northern Hemisphere."
    )

    add_body_text(doc,
        "To distinguish biological from institutional drivers, we exploited a natural "
        "experiment: geographic variation in academic/fiscal year calendars. Countries were "
        "grouped by academic year start month: Japan and South Korea (March/April start, "
        "n=250), USA/Europe/China (September/October start, n=2,825), and "
        "Australia/New Zealand/Brazil (January/February start, Southern Hemisphere, n=83). "
        "Using Pearson's \u03c7\u00b2 and circular Rayleigh tests (Mardia and Jupp, 2000), we found "
        "a striking dissociation (Figure 1)."
    )

    add_body_text(doc,
        "Only the Japan/South Korea group showed significant seasonality (\u03c7\u00b2=32.8, df=11, "
        "p=0.0006, Rayleigh R=0.172) with a March peak perfectly aligned with Japan's "
        "fiscal year-end. The much larger September-start group (n=2,825) showed no "
        "significant seasonal pattern (\u03c7\u00b2=15.0, p=0.13). A sensitivity analysis confirmed "
        "that Japan\u2014just 3.3% of the total dataset\u2014drove the overall signal: excluding "
        "Japan shifted the peak from March to January and, when restricted to records with "
        "known country affiliation (n=3,195), eliminated statistical significance entirely "
        "(p=0.128)."
    )

    add_body_text(doc,
        "This result carries a critical methodological lesson. What appeared to be evidence "
        "of seasonal biological influence on stem cell research was in fact an artifact of "
        "Japan's March 31 fiscal year-end, which creates a rush to submit manuscripts and "
        "deposit data before annual reporting deadlines. Japan contributed only 3.3% of "
        "total datasets (200/6,101) but accounted for 15% of all March submissions\u2014a "
        "nearly two-fold over-representation relative to its overall share. The fiscal "
        "year-end effect is well-documented in Japanese academic culture: grant reports "
        "are due by March 31, creating institutional pressure to finalize data and "
        "publications before the deadline."
    )

    add_body_text(doc,
        "An exploratory analysis of the relationship between PSC dataset volume and solar "
        "activity indices (NOAA F10.7 flux) illustrates a complementary pitfall. The raw "
        "correlation was r=0.48, suggesting an intriguing link\u2014until detrending removed "
        "the common upward trajectory of both iPSC research output (exponential growth "
        "since 2006) and Solar Cycle 24/25 ascent, yielding a residual correlation of "
        "r=0.02 (p=0.75). Co-trending time series frequently produce spurious correlations "
        "that evaporate once the shared secular trend is removed."
    )

    add_body_text(doc,
        "Geographic analysis further constrained interpretation: 97% of GEO datasets with "
        "identifiable country affiliations originated from Northern Hemisphere institutions, "
        "rendering hemisphere-inversion tests\u2014the most powerful diagnostic for "
        "distinguishing solar-driven from institutional effects\u2014impractical with GEO data "
        "alone. The broader implication extends beyond stem cell biology. Any field that "
        "relies on publication or data-deposition timing as a proxy for experimental timing "
        "risks confounding institutional rhythms with biological ones. GEO submission dates "
        "reflect a complex convolution of experiment timing, analysis duration, writing, "
        "peer review, and institutional incentives\u2014with a typical lag of 6\u201318 months from "
        "bench experiment to data deposition. Our natural experiment demonstrates that "
        "public database metadata cannot reliably distinguish biological seasonality from "
        "institutional rhythms. This finding should serve as a cautionary note for "
        "meta-research studies that use submission timestamps as temporal proxies."
    )

    # ──────────────────────────────────────────────
    # Section 6: The Path Forward
    # ──────────────────────────────────────────────
    add_heading(doc, "The Path Forward: Prospective Environmental Monitoring", level=2)

    add_body_text(doc,
        "Our GEO analysis demonstrates that retrospective database mining is insufficient "
        "to detect or rule out environmental seasonality in stem cell differentiation. The "
        "IVF evidence (Leathersich et al., 2023; Wang et al., 1992) suggests that such "
        "effects exist in early-stage cells, and the vulnerability evidence "
        "(Panina et al., 2021; McCreery et al., 2025) suggests PSCs are plausible targets. "
        "What is needed is prospective, controlled environmental monitoring paired with "
        "differentiation outcome tracking."
    )

    add_body_text(doc,
        "We propose a three-phase approach. Phase I: equip PSC facilities with continuous "
        "multi-parameter sensors (humidity, illuminance, ELF-EMF, barometric pressure, VOC "
        "levels, vibration) at 1-minute resolution alongside routine differentiation "
        "outcomes for \u226512 months. Modern IoT sensor packages achieve this at modest cost "
        "(<$5,000 per laboratory). Phase II: analyze existing data from IVF registries "
        "(HFEA in the UK, ANZARD in Australia) that contain date-stamped cycle outcomes with "
        "latitude information\u2014these can directly test the hemisphere-inversion prediction "
        "for photoperiod-driven effects. Phase III: for variables identified in Phases I\u2013II "
        "as significantly correlated with outcomes, conduct prospective controlled "
        "interventions (e.g., humidity-controlled vs. uncontrolled incubators, light-tight "
        "hoods vs. standard biosafety cabinets, magnetic shielding)."
    )

    add_body_text(doc,
        "The hemisphere-inversion test provides a powerful diagnostic criterion. Variables "
        "whose effects are phase-inverted between Northern and Southern Hemisphere "
        "laboratories (e.g., summer advantage in both Perth and Boston, six months apart) "
        "implicate solar-driven mechanisms\u2014photoperiod, temperature, UV flux. Variables "
        "whose effects are hemisphere-synchronous (e.g., identical timing in both "
        "hemispheres) implicate geomagnetic or cosmic ray pathways, which are global rather "
        "than latitude-dependent. A coordinated multi-center study spanning both hemispheres "
        "could decompose seasonal variation into these mechanistic categories."
    )

    add_body_text(doc,
        "The key principle is that monitoring imposes no change on existing workflows\u2014it "
        "simply makes the invisible visible. Unlike prospective intervention studies, which "
        "require protocol modifications and ethical approvals, passive environmental "
        "monitoring can be implemented immediately in any laboratory. The resulting datasets, "
        "when combined across multiple sites and latitudes, would constitute a resource of "
        "enormous value for the field\u2014analogous to weather station networks that transformed "
        "meteorology from an anecdotal practice to a predictive science."
    )

    add_body_text(doc,
        "The technical requirements for Phase I monitoring are straightforward. "
        "A minimum sensor array would include: (1) temperature/humidity loggers (\u00b10.1\u00b0C, "
        "\u00b11% RH) inside incubators and at bench level; (2) a photodiode sensor in the "
        "biosafety cabinet recording cumulative lux-hours of exposure per manipulation; "
        "(3) a VOC sensor (ppb-level total VOC) with seasonal baseline tracking; (4) an "
        "ELF-EMF sensor (3-axis, nT resolution) placed adjacent to incubators; and (5) a "
        "barometric pressure logger. Modern IoT platforms can integrate all five data "
        "streams with automatic cloud upload, enabling multi-site meta-analyses without "
        "manual data wrangling. The critical companion requirement is systematic recording "
        "of differentiation outcomes per batch\u2014ideally including not just binary success/"
        "failure but quantitative markers (flow cytometry percentages, gene expression "
        "levels, electrophysiological parameters) that permit correlation analysis."
    )

    # Figure 2 reference
    add_body_text(doc,
        "Figure 2 presents a research roadmap organized by the expected hemisphere-"
        "dependence of each variable and the current level of evidence."
    )

    # ──────────────────────────────────────────────
    # Section 7: Implications
    # ──────────────────────────────────────────────
    add_heading(doc, "Implications for Cell Therapy Manufacturing", level=2)

    add_body_text(doc,
        "As PSC-derived cell therapies advance toward clinical application "
        "(Kirkeby et al., 2025), manufacturing consistency becomes a regulatory imperative. "
        "Current Good Manufacturing Practice (cGMP) facilities control temperature, "
        "humidity, and particulates, but do not routinely monitor illuminance, EMF, VOCs, "
        "or barometric pressure. If any of these variables meaningfully affect "
        "differentiation, current manufacturing harbors hidden batch-to-batch variability "
        "attributed to biological stochasticity but potentially environmental and "
        "correctable. The economic argument is compelling: failed batches in autologous cell "
        "therapy represent lost patient tissue and treatment delays. Even modest improvements "
        "in first-pass success rates through environmental optimization could translate into "
        "significant cost savings and reduced time-to-treatment."
    )

    add_body_text(doc,
        "The IVF field's experience is instructive: when laboratories invested in VOC "
        "filtration and positive-pressure air handling, blastocyst rates increased by 18% "
        "and live births by 8% (Agarwal et al., 2017). These improvements were achieved not "
        "through better biological protocols but through better environmental control\u2014"
        "precisely the paradigm shift we advocate for PSC manufacturing. The convergence of "
        "clinical-grade cell therapy production with modern IoT environmental monitoring "
        "capabilities creates an opportune moment to establish new standards for "
        "'environmentally aware' cell culture."
    )

    # ──────────────────────────────────────────────
    # Section 8: What This Opinion Does Not Argue
    # ──────────────────────────────────────────────
    add_heading(doc, "What This Opinion Does Not Argue", level=2)

    add_body_text(doc,
        "We wish to be explicit about the boundaries of our argument. We are not claiming "
        "that seasonal photoperiod directly drives PSC differentiation outcomes\u2014our GEO "
        "analysis demonstrates precisely the opposite: the apparent seasonal signal in "
        "public databases is institutional, not biological. We are not claiming that "
        "environmental variables are the dominant source of differentiation variability\u2014"
        "genetic background, passage number, and operator technique remain important. "
        "Rather, we argue that environmental variables constitute an additional, correctable "
        "source of variation that has been systematically overlooked because it is never "
        "measured. The null hypothesis\u2014that variables like humidity, VOCs, and ambient light "
        "have zero effect on PSC differentiation\u2014is contradicted by the evidence reviewed "
        "above (Panina et al., 2021; McCreery et al., 2025; Bi et al., 2020), but the "
        "magnitude of effects under real-world laboratory conditions remains unknown. Only "
        "prospective measurement can resolve this question."
    )

    # ──────────────────────────────────────────────
    # Conclusion
    # ──────────────────────────────────────────────
    add_heading(doc, "Conclusion", level=2)

    add_body_text(doc,
        "We have presented three convergent arguments. First, PSCs are uniquely vulnerable "
        "to environmental perturbations at magnitudes that leave differentiated cells "
        "unaffected\u2014from osmotic stress (McCreery et al., 2025) to airborne pollutants "
        "(Bi et al., 2020) to circadian disruption (Sato et al., 2023). Second, multiple "
        "environmental variables in typical PSC laboratories remain unmonitored and exhibit "
        "seasonal variation (Mizuno et al., 2020). Third, our natural experiment using GEO "
        "metadata demonstrates that apparent seasonality in stem cell databases is an "
        "institutional calendar artifact, meaning that database-mining approaches cannot "
        "resolve the question\u2014only prospective measurement can."
    )

    add_body_text(doc,
        "More broadly, this work exposes a methodological complacency inherent in "
        "clonal research. The assumption that genetically homogeneous systems require no "
        "confounding adjustment has allowed non-genetic sources of variation\u2014environmental, "
        "institutional, and temporal\u2014to accumulate undetected. Our GEO natural experiment "
        "demonstrates how a seemingly homogeneous global dataset can be driven by a single "
        "country's fiscal calendar. This lesson extends beyond PSC biology to any field "
        "where clonal or isogenic systems create a false sense of experimental control: "
        "genetic uniformity does not guarantee environmental uniformity, and aggregate "
        "data from clonal experiments are not immune to confounding."
    )

    add_body_text(doc,
        "The history of experimental biology is punctuated by discoveries that 'noise' was "
        "in fact signal from an unmeasured variable. The stem cell field may be at such an "
        "inflection point. We have spent two decades optimizing recipes while ignoring the "
        "kitchen. The convergence of evidence from IVF seasonality "
        "(Leathersich et al., 2023; Wang et al., 1992), circadian regulation of "
        "pluripotency (Sato et al., 2023; Ameneiro et al., 2020), and environmental "
        "monitoring of cell-processing facilities (Mizuno et al., 2020) suggests that the "
        "invisible variables are neither negligible nor intractable. They merely need to be "
        "measured."
    )

    add_body_text(doc,
        "The practical path forward is clear. First, equip PSC laboratories with "
        "multi-parameter environmental sensors and record outcomes systematically for "
        "12 months. Second, analyze IVF registry data (HFEA, ANZARD) for "
        "hemisphere-dependent seasonal effects on embryo quality. Third, for any "
        "variable identified as significant, conduct controlled interventions. The "
        "cost of environmental monitoring is negligible compared to the cost of failed "
        "differentiation batches in clinical manufacturing (Kirkeby et al., 2025; "
        "Karagiannis et al., 2019). If even one invisible variable proves "
        "consequential\u2014as VOCs proved consequential in IVF (Agarwal et al., 2017)\u2014the "
        "return on investment for the field will be transformative."
    )

    # The conclusion's final paragraph should summarize themes and identify
    # unanswered questions (per HR guidelines)
    add_body_text(doc,
        "In summary, this Opinion has identified a critical gap in stem cell research: "
        "the systematic neglect of environmental variables in PSC culture. Key unanswered "
        "questions include: (i) which specific environmental variables (humidity, VOCs, "
        "light, EMF, barometric pressure) have the largest effect on differentiation "
        "outcomes? (ii) Are these effects consistent across cell lines and differentiation "
        "protocols? (iii) Can the hemisphere-inversion test discriminate solar-driven from "
        "geomagnetic mechanisms using IVF registry data? (iv) What are the minimum "
        "monitoring requirements for clinically meaningful environmental quality control? "
        "Addressing these questions through prospective, multi-site environmental profiling "
        "would represent a paradigm shift from recipe optimization to kitchen optimization "
        "in stem cell biology."
    )

    # ──────────────────────────────────────────────
    # Data and Code Availability
    # ──────────────────────────────────────────────
    add_heading(doc, "Data Availability", level=2)
    add_body_text(doc,
        "The GEO metadata analysis used publicly available data accessed through the NCBI "
        "GEO API (Barrett et al., 2013) (search query: 'iPSC OR ESC differentiation'; "
        "date range: 2001\u20132026). Solar activity indices were obtained from the NOAA Space "
        "Weather Prediction Center. Analysis code, processed datasets, and "
        "figure-generation scripts are available at "
        "[GitHub repository URL to be inserted upon acceptance]."
    )

    # ──────────────────────────────────────────────
    # Acknowledgments
    # ──────────────────────────────────────────────
    add_heading(doc, "Acknowledgements", level=2)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run("[To be added]")
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # ──────────────────────────────────────────────
    # Authors' Roles
    # ──────────────────────────────────────────────
    add_heading(doc, "Authors' Roles", level=2)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run("[To be added — e.g., T.O. conceived the study, performed the GEO "
                     "analysis, and drafted the manuscript. All authors reviewed and "
                     "approved the final version.]")
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # ──────────────────────────────────────────────
    # Funding
    # ──────────────────────────────────────────────
    add_heading(doc, "Funding", level=2)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run("[To be added — or 'No external funding was received for this work.']")
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # ──────────────────────────────────────────────
    # Conflict of Interest
    # ──────────────────────────────────────────────
    add_heading(doc, "Conflict of Interest", level=2)
    add_body_text(doc, "The authors declare no competing interests.")

    # ──────────────────────────────────────────────
    # Figure Legends (at end of manuscript, per HR guidelines)
    # ──────────────────────────────────────────────
    add_heading(doc, "Figure Legends", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run("Figure 1. Academic calendar variation as a natural experiment. ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run(
        "(A) Normalized monthly distributions of GEO PSC dataset submissions by academic "
        "year group. The March/April-start group (Japan/Korea, red) shows a pronounced "
        "March peak absent in the September-start group (USA/Europe/China, blue). "
        "(B) Country-level heatmap: the pattern is driven predominantly by Japan. "
        "(C) Circular mean directions (Rayleigh vectors) for each group. "
        "(D) Statistical summary. Only Japan/Korea achieves significance (p<0.001). "
        "n=6,101 total datasets; 3,195 with country affiliation."
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0, space_before=Pt(12))
    run = p.add_run(
        "Figure 2. Research roadmap for environmental profiling of PSC culture. "
    )
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run(
        "(A) Evidence matrix: uncontrolled environmental variables classified by "
        "hemisphere-dependence (inverted = solar-driven; synchronous = geomagnetic) and "
        "evidence level (direct PSC, indirect IVF/cell culture, theoretical). "
        "(B) Three-phase investigation strategy from passive monitoring through "
        "retrospective data mining to controlled intervention. "
        "Variables whose effects are hemisphere-inverted implicate solar/photoperiod "
        "mechanisms; hemisphere-synchronous effects implicate geomagnetic pathways."
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Alt text for figures (HR requirement)
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0, space_before=Pt(12))
    run = p.add_run("Alt text for Figure 1: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run(
        "Four-panel figure showing monthly distribution of GEO PSC dataset submissions "
        "grouped by academic calendar region. Panel A shows normalized bar charts, Panel B "
        "a country-level heatmap, Panel C circular statistics vectors, and Panel D a "
        "statistical summary table. The March peak is driven by Japan's fiscal year-end."
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0, space_before=Pt(6))
    run = p.add_run("Alt text for Figure 2: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run(
        "Two-panel research roadmap figure. Panel A shows an evidence matrix classifying "
        "environmental variables by hemisphere-dependence and evidence level. Panel B shows "
        "a three-phase investigation strategy from passive monitoring to controlled intervention."
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # ──────────────────────────────────────────────
    # References (alphabetical, author-date format)
    # ──────────────────────────────────────────────
    add_heading(doc, "References", level=2)

    sorted_refs = get_sorted_references()
    for ref in sorted_refs:
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=Pt(6), line_spacing=2.0)
        ref_text = format_ref_for_list(ref)
        run = p.add_run(ref_text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    # Save
    out_path = os.path.join(OUTPUT_DIR, "HumReprod_Opinion_InvisibleVariables.docx")
    doc.save(out_path)
    print(f"Manuscript saved to: {out_path}")

    # Word count (body text only, excluding references/legends/metadata)
    word_count = 0
    in_body = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "The Reproducibility Crisis Has an Environmental Blind Spot":
            in_body = True
            continue
        if text in ("Data Availability", "Acknowledgements", "Authors' Roles",
                    "Funding", "Conflict of Interest", "Figure Legends", "References"):
            in_body = False
            continue
        if in_body and text:
            word_count += len(text.split())
    print(f"Approximate body word count (excl. refs/legends/metadata): {word_count}")

    return out_path


if __name__ == "__main__":
    create_manuscript()
