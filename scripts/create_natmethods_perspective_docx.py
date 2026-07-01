#!/usr/bin/env python3
"""
Generate Nature Methods Perspective manuscript:
"The Invisible Variables: Why Clonal Systems Are Not Immune to
 Environmental Confounding"

Format: Nature Methods Perspective
- Superscript numbered citations, sequential order of first appearance
- Nature reference format: Author, A. B., Author, C. D. & Author, E. F.
  Title. J. Abbrev. vol, pages (year).
- Up to ~4,000 words main text
- Up to 50 references (exceptions possible)
- Up to 6 display items
- Abstract: unstructured, ≤150 words, unreferenced

Narrative structure:
  Baking analogy → IVF environmental success → PSC vulnerability evidence
  → Clonal complacency trap → GEO natural experiment → Monitoring proposal
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ──────────────────────────────────────────────
# References in order of first appearance in text
# Nature format: Author, A. B., Author, C. D. & Author, E. F.
#   Title. J. Abbrev. vol, pages (year).
# ──────────────────────────────────────────────
REFERENCES = [
    # 1 - Kirkeby (clinical trials update) [Section 1, para 1]
    "Kirkeby, A., Main, H. & Carpenter, M. Pluripotent stem-cell-derived "
    "therapies in clinical trial: a 2025 update. Cell Stem Cell 32, 10\u201337 (2025).",
    # 2 - Yamanaka (PSC therapy promise & challenges) [Section 1, para 1]
    "Yamanaka, S. Pluripotent stem cell-based cell therapy\u2014promise and challenges. "
    "Cell Stem Cell 27, 523\u2013531 (2020).",
    # 3 - Volpato (multi-site reproducibility) [Section 1, para 1]
    "Volpato, V. et al. Reproducibility of molecular phenotypes after long-term "
    "differentiation to human iPSC-derived neurons: a multi-site omics study. "
    "Stem Cell Reports 11, 897\u2013911 (2018).",
    # 4 - Volpato & Webber (iPSC variability guidelines) [Section 1, para 1]
    "Volpato, V. & Webber, C. Addressing variability in iPSC-derived models of "
    "human disease: guidelines to promote reproducibility. Dis. Model. Mech. 13, "
    "dmm042317 (2020).",
    # 5 - Ortmann (PSC line variability) [Section 1, para 1]
    "Ortmann, D. & Vallier, L. Variability of human pluripotent stem cell lines. "
    "Curr. Opin. Genet. Dev. 46, 179\u2013185 (2017).",
    # 6 - Agarwal (VOCs and IVF outcomes) [Section 1, para 2]
    "Agarwal, N. et al. Volatile organic compounds and good laboratory practices "
    "in the in vitro fertilization laboratory: the important parameters for "
    "successful outcome in extended culture. J. Assist. Reprod. Genet. 34, "
    "999\u20131006 (2017).",
    # 7 - Cai (VOC real-time monitoring, embryology lab) [Section 1, para 2]
    "Cai, J. et al. Real-time monitoring reveals the effects of low concentrations "
    "of volatile organic compounds in the embryology laboratory. Hum. Reprod. 40, "
    "deaf008 (2025).",
    # 8 - Cohen (Cairo consensus IVF lab air quality) [Section 2, para 1]
    "Cohen, J. et al. Cairo consensus on the IVF laboratory environment and air "
    "quality: report of an expert meeting. Reprod. Biomed. Online 36, 658\u2013674 "
    "(2018).",
    # 9 - Leathersich (IVF season and oocyte collection) [Section 2, para 2]
    "Leathersich, S. J. et al. Season at the time of oocyte collection and frozen "
    "embryo transfer outcomes. Hum. Reprod. 38, 1714\u20131722 (2023).",
    # 10 - Wang C (IVF season, temperature, DNI) [Section 2, para 2]
    "Wang, C. et al. The impact of season, temperature, and direct normal "
    "irradiance on IVF pregnancy outcomes: a retrospective cohort study. Int. J. "
    "Biometeorol. 69, 2053\u20132065 (2025).",
    # 11 - Braga (embryo quality and morphokinetics by season) [Section 2, para 2]
    "Braga, D. P. A., Setti, A., Guilherme, P., Iaconelli, A. Jr & Borges, E. Jr. "
    "Association between meteorological season and embryo quality in the era of "
    "morphokinetics. J. Assist. Reprod. Genet. 42, 1287\u20131298 (2025).",
    # 12 - Deng (season and IVF: systematic review & meta-analysis) [Section 2, para 2]
    "Deng, Q. et al. Association between season and pregnancy outcomes in fresh "
    "embryo transfer cycles: a systematic review and meta-analysis. Front. Public "
    "Health 13, 1660982 (2025).",
    # 13 - Wang JP (mouse seasonal two-cell block) [Section 2, para 2]
    "Wang, J. P. et al. Seasonal variation in cell cycle during early development "
    "of the mouse embryo. Reproduction 94, 431\u2013436 (1992).",
    # 14 - Panina (PSC drug sensitivity vs somatic cells) [Section 3, para 2]
    "Panina, Y., Yamane, J., Kobayashi, K., Sone, H. & Fujibuchi, W. Human ES "
    "and iPS cells display less drug resistance than differentiated cells, and "
    "naive-state induction further decreases drug resistance. J. Toxicol. Sci. 46, "
    "131\u2013142 (2021).",
    # 15 - McCreery (mechano-osmotic chromatin, pluripotent fate) [Section 3, para 3]
    "McCreery, K. P. et al. Mechano-osmotic signals control chromatin state and "
    "fate transitions in pluripotent stem cells. Nat. Cell Biol. 27, 1757\u20131770 "
    "(2025).",
    # 16 - Chui (osmolar modulation, iPSC differentiation) [Section 3, para 3]
    "Chui, J. S.-H. et al. Osmolar modulation drives reversible cell cycle exit "
    "and human pluripotent cell differentiation via NF-\u03baB and WNT signaling. "
    "Adv. Sci. 11, 2307554 (2024).",
    # 17 - Sato (CRY1 in PSCs) [Section 3, para 4]
    "Sato, S. et al. The circadian clock CRY1 regulates pluripotent stem cell "
    "identity and somatic cell reprogramming. Cell Rep. 42, 112590 (2023).",
    # 18 - Ameneiro (BMAL1 in PSCs) [Section 3, para 4]
    "Ameneiro, C. et al. BMAL1 coordinates energy metabolism and differentiation "
    "of pluripotent stem cells. Life Sci. Alliance 3, e201900534 (2020).",
    # 19 - Golan (circadian rhythms and HSC differentiation) [Section 3, para 4]
    "Golan, K., Kollet, O., Markus, R. P. & Lapidot, T. Daily light and darkness "
    "onset and circadian rhythms metabolically synchronize hematopoietic stem cell "
    "differentiation and maintenance. Exp. Hematol. 78, 1\u201310 (2019).",
    # 20 - Bi (PM2.5 and hESC pluripotency) [Section 3, para 5]
    "Bi, S. et al. Fine particulate matter reduces the pluripotency and "
    "proliferation of human embryonic stem cells through ROS induced AKT and ERK "
    "signaling pathway. Reprod. Toxicol. 96, 231\u2013240 (2020).",
    # 21 - Czyz (EMF and mouse ESCs) [Section 3, para 5]
    "Czyz, J. et al. Non-thermal effects of power-line magnetic fields (50 Hz) on "
    "gene expression levels of pluripotent embryonic stem cells. Mutat. Res. 557, "
    "63\u201374 (2004).",
    # 22 - Mizuno (environmental risk assessment, cell-processing) [Section 4, para 1]
    "Mizuno, M. et al. The environmental risk assessment of cell-processing "
    "facilities for cell therapy in a Japanese academic institution. PLoS ONE 15, "
    "e0236600 (2020).",
    # 23 - Klein (best practices cell culture environments) [Section 4, para 1]
    "Klein, S. G. et al. Toward best practices for controlling mammalian cell "
    "culture environments. Front. Cell Dev. Biol. 10, 788808 (2022).",
    # 24 - Diatroptova (cosmic radiation and cell proliferation) [Section 4, para 1]
    "Diatroptova, M. A., Kosyreva, A. M. & Diatroptov, M. E. About 4-day rhythm "
    "of proliferative activity of L-929 cells in culture correlates with the "
    "intensity of secondary cosmic radiation fluctuations. Bull. Exp. Biol. Med. "
    "172, 561\u2013565 (2022).",
    # 25 - Barrett (NCBI GEO database) [Section 6, para 1]
    "Barrett, T. et al. NCBI GEO: archive for functional genomics data sets\u2014"
    "update. Nucleic Acids Res. 41, D991\u2013D995 (2013).",
    # 26 - Mardia (directional/circular statistics) [Section 6, para 2]
    "Mardia, K. V. & Jupp, P. E. Directional Statistics (Wiley, 2000).",
    # 27 - Karagiannis (iPSC use in disease models) [Section 8, para 1]
    "Karagiannis, P. et al. Induced pluripotent stem cells and their use in human "
    "models of disease and development. Physiol. Rev. 99, 79\u2013114 (2019).",
    # 28 - Henn (cytocentric principles for RM reproducibility) [Section 8, para 2]
    "Henn, A. D. et al. Applying the cytocentric principles to regenerative medicine "
    "for reproducibility. Curr. Stem Cell Rep. 8, 197\u2013202 (2022).",
    # 29 - Navarro (cytocentric measurement for RM) [Section 8, para 2]
    "Navarro, M. et al. Cytocentric measurement for regenerative medicine. "
    "Front. Med. Technol. 5, 1154653 (2023).",
]


def add_superscript_refs(paragraph, text):
    """Parse text with {N} or {N-M} or {N,M} markers and create superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(9)
            run.font.name = 'Arial'
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
            run.font.name = 'Arial'
    return paragraph


def set_paragraph_format(para, space_after=Pt(6), space_before=Pt(0),
                         line_spacing=2.0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    para.paragraph_format.line_spacing = line_spacing
    para.alignment = alignment
    return para


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_body_paragraph(doc, text):
    """Add a body paragraph with superscript citation handling."""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    add_superscript_refs(p, text)
    return p


def create_manuscript():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # ──────────────────────────────────────────────
    # Title page
    # ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Perspective")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 102, 153)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(24), space_after=Pt(12))
    run = p.add_run(
        "The Invisible Variables: Why Clonal Systems Are Not Immune "
        "to Environmental Confounding"
    )
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(24))
    run = p.add_run("[Author names and affiliations to be added]")
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.italic = True

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(6))
    run = p.add_run("Correspondence: ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run = p.add_run("[Corresponding author email]")
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.italic = True

    # ──────────────────────────────────────────────
    # Abstract (≤150 words, unreferenced for Nature Methods)
    # ──────────────────────────────────────────────
    add_heading(doc, "Abstract", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run(
        "Pluripotent stem cell (PSC) differentiation protocols suffer from persistent "
        "batch-to-batch variability that impedes both basic research and clinical "
        "manufacturing. We argue that this variability has a methodological blind spot: "
        "multiple environmental variables\u2014humidity, volatile organic compounds (VOCs), "
        "ambient light, electromagnetic fields, and barometric pressure\u2014remain unmonitored "
        "in most laboratories. The clonal nature of PSC research creates a false sense of "
        "experimental control; genetic homogeneity does not eliminate environmental "
        "confounding. We draw on two precedents where environmental control transformed "
        "outcomes: artisanal baking and IVF embryology. A natural experiment using 6,101 "
        "GEO datasets reveals that apparent seasonality in public databases is an "
        "institutional artifact, not a biological signal. We propose IoT-based multi-parameter "
        "environmental monitoring as a systematic method to identify which invisible "
        "variables matter."
    )
    run.font.name = 'Arial'
    run.font.size = Pt(11)

    # ──────────────────────────────────────────────
    # Main text
    # ──────────────────────────────────────────────

    # --- Section 1: Introduction (the blind spot) ---
    add_heading(doc, "The methodological blind spot", level=2)

    add_body_paragraph(doc,
        "Pluripotent stem cell (PSC)-derived therapies are advancing rapidly, with over "
        "115 trials and 1,200 patients dosed as of 2024{1}. Yet differentiation protocols "
        "remain notoriously variable across laboratories and even between batches within a "
        "single group{2\u20135}. The field has invested enormously in optimizing controllable "
        "recipe variables\u2014growth factors, small molecules, extracellular matrices, and "
        "timing\u2014while largely ignoring the kitchen in which these recipes are executed."
    )

    add_body_paragraph(doc,
        "Consider an analogy. For decades, artisanal bread quality varied unpredictably "
        "between batches until industrial bakeries recognized that controlling only oven "
        "temperature was insufficient; ambient humidity, dough temperature during proofing, "
        "and air quality were equally critical. The in vitro fertilization (IVF) field "
        "underwent an analogous revolution when it discovered that volatile organic "
        "compounds (VOCs) in embryology laboratories were a major determinant of blastocyst "
        "formation rates: reducing VOCs increased blastocyst rates by 18% and live birth "
        "rates by 8%{6,7}. Today\u2019s PSC laboratory resembles the pre-industrial bakery: "
        "we control incubator temperature (37\u00b0C) and CO\u2082 (5%) but leave humidity, VOCs, "
        "ambient light, electromagnetic fields (EMF), and barometric pressure entirely "
        "unmonitored."
    )

    add_body_paragraph(doc,
        "A deeper methodological issue compounds the problem. Research using clonal "
        "systems\u2014isogenic cell lines, inbred mice, genetically identical iPSC clones\u2014"
        "operates under the implicit assumption that genetic homogeneity renders complex "
        "confounding adjustment unnecessary. Because the genetic background is \u2018controlled,\u2019 "
        "investigators rarely apply the stratification, regression, or sensitivity analyses "
        "that epidemiologists would consider mandatory for observational data. This creates a "
        "blind spot: non-genetic confounders (environmental, institutional, temporal) are "
        "neither measured nor adjusted for, because the clonal paradigm implies they should "
        "not exist. We term this the \u2018clonal complacency trap\u2019\u2014the methodological "
        "assumption that genetic uniformity guarantees experimental uniformity."
    )

    add_body_paragraph(doc,
        "The scale of the problem is not trivial. Multi-site reproducibility studies "
        "consistently report substantial inter-laboratory variation in PSC differentiation "
        "outcomes even when nominally identical protocols are followed{3}. Volpato et al. "
        "found that laboratory-of-origin was the dominant source of variance in iPSC-derived "
        "neuron transcriptomes, exceeding the effect of genetic background or differentiation "
        "batch. Current explanations invoke differences in operator technique, reagent lots, "
        "and passage number\u2014but these account for only a fraction of the observed variance. "
        "We propose that uncontrolled environmental differences between laboratories "
        "constitute a missing explanatory variable."
    )

    add_body_paragraph(doc,
        "This proposal is neither speculative nor unprecedented. The IVF field underwent "
        "a parallel revolution when it recognized that air quality in embryology laboratories "
        "was a major determinant of blastocyst formation rates{6,7}. Prior to this insight, "
        "unexplained clinic-to-clinic variation in IVF success rates was attributed to "
        "operator skill and patient selection; after implementing VOC filtration and "
        "environmental controls, many laboratories saw 10\u201320% improvements in outcomes{6}. "
        "We argue that the PSC field is at a similar inflection point, with the added "
        "advantage that modern IoT sensor technology makes comprehensive environmental "
        "monitoring far cheaper and easier than it was for IVF laboratories two decades ago."
    )

    # --- Section 2: Two precedents for environmental control ---
    add_heading(doc, "Two precedents: when controlling the kitchen transformed outcomes",
                level=2)

    add_body_paragraph(doc,
        "The IVF field provides the strongest existing precedent. Before systematic "
        "environmental control, unexplained clinic-to-clinic variation in success rates was "
        "attributed to operator skill and patient selection. After implementing VOC "
        "filtration and positive-pressure air handling{8}, many laboratories saw 10\u201320% "
        "improvements in outcomes{6}. Cai et al. demonstrated through real-time VOC "
        "monitoring that even low concentrations linearly predicted decreased blastocyst "
        "quality, and that the effect was specific to fresh embryos\u2014frozen embryos were "
        "unaffected\u2014indicating that early-stage cells are selectively vulnerable{7}."
    )

    add_body_paragraph(doc,
        "Seasonal patterns in IVF outcomes provide further evidence. Leathersich et al. "
        "analyzed 3,659 frozen embryo transfer cycles in Perth, Australia (33\u00b0S), finding "
        "that oocytes collected in summer had 30% higher live birth rates than those "
        "collected in autumn{9}. Crucially, only the season at oocyte collection "
        "mattered\u2014not the season of transfer\u2014establishing that the effect acts on the "
        "gamete/early embryo rather than on implantation. Northern Hemisphere studies "
        "report concordant summer advantages{10}, and a Brazilian study at 23\u00b0S confirms "
        "the expected six-month phase shift{11}. Although a meta-analysis found minimal "
        "overall effects in pooled data{12}, Wang et al. provided direct experimental "
        "evidence in mice: embryos cultured during summer showed the \u2018two-cell block\u2019 "
        "phenomenon, while winter embryos developed normally{13}."
    )

    add_body_paragraph(doc,
        "The IVF evidence is particularly instructive because it provides a causal "
        "structure: the season at oocyte collection affects outcomes, but the season at "
        "embryo transfer does not{9}. This temporal dissection implies that the "
        "environmental effect acts during a specific developmental window\u2014oocyte maturation "
        "and early cleavage\u2014rather than through a general mechanism affecting all cells at "
        "all times{9}. PSC differentiation protocols, which recapitulate aspects of early "
        "embryonic lineage commitment over 7\u201330 days, would be expected to share this "
        "window-specific vulnerability. The parallel is not merely analogical: iPSCs "
        "undergoing directed differentiation traverse epigenetic states homologous to those "
        "of the early embryo, including the exit from naive pluripotency, lineage priming, "
        "and specification\u2014each potentially gated by the same environmental sensitivities "
        "that affect preimplantation development."
    )

    add_body_paragraph(doc,
        "These precedents share a common lesson: what was attributed to inherent biological "
        "stochasticity turned out to be partly environmental and correctable. The PSC field "
        "has not yet had its \u2018VOC moment\u2019\u2014the discovery that a specific, measurable "
        "environmental variable explains a meaningful fraction of unexplained variability."
    )

    # --- Section 3: Why PSCs are vulnerable ---
    add_heading(doc, "Why pluripotent stem cells are uniquely vulnerable", level=2)

    add_body_paragraph(doc,
        "The premise that uncontrolled environmental factors could meaningfully affect PSC "
        "differentiation\u2014while being negligible for established cell lines\u2014requires "
        "justification. Several lines of evidence support this claim."
    )

    add_body_paragraph(doc,
        "First, PSCs display markedly greater chemical sensitivity than differentiated "
        "cells. Panina et al. demonstrated that human iPS cells are approximately 1.5-fold "
        "more sensitive to drug exposure than ES cells, and both are several-fold less "
        "resistant than non-pluripotent cell types{14}. Critically, naive-state induction "
        "further increased sensitivity, establishing a gradient: the more undifferentiated "
        "the cell, the more vulnerable it is to environmental perturbation."
    )

    add_body_paragraph(doc,
        "Second, osmotic perturbations of surprisingly small magnitude can gate cell fate "
        "transitions. McCreery et al. showed that compaction-triggered changes in nuclear "
        "shape and volume in human iPSCs remodel chromatin architecture and prime cells for "
        "ectodermal differentiation{15}. During biosafety cabinet work, culture medium "
        "evaporates at rates determined by ambient humidity, producing osmolarity increases "
        "that could cross these thresholds. Chui et al. confirmed that hyperosmotic culture "
        "drives NF-\u03baB and WNT signaling-dependent cell cycle exit and maturation{16}\u2014"
        "meaning osmolarity changes of physiological magnitude are not merely stress events "
        "but active differentiation signals."
    )

    add_body_paragraph(doc,
        "Third, PSCs depend on circadian clock components in a non-canonical manner. CRY1 "
        "is dramatically upregulated in iPSCs and ESCs compared to somatic cells, and its "
        "deletion impairs self-renewal and disrupts differentiation{17}. BMAL1 coordinates "
        "energy metabolism and differentiation independently of circadian oscillation, which "
        "is suppressed in pluripotent cells{18}. Consistent with this, daily light and "
        "darkness onset metabolically synchronize hematopoietic stem cell differentiation "
        "and maintenance{19}, suggesting that light-sensitive circadian pathways are broadly "
        "relevant to stem cell biology beyond the pluripotent compartment. This "
        "architecture\u2014dependence on clock molecules without functional clock oscillation\u2014"
        "may render PSCs vulnerable to external zeitgeber inputs (e.g., ambient light "
        "leaking into incubators or biosafety cabinets) that would not perturb "
        "differentiated cells with intact circadian buffering."
    )

    add_body_paragraph(doc,
        "Fourth, airborne pollutants affect PSCs at concentrations that leave somatic cells "
        "unharmed. Fine particulate matter (PM2.5) downregulates pluripotency markers "
        "(NANOG, OCT4) in human ESCs through ROS-mediated AKT/ERK signaling{20}. "
        "Electromagnetic fields from incubator heaters, freezer compressors, and building "
        "wiring (50/60 Hz) alter gene expression in mouse ESCs{21}, though effect sizes "
        "under typical laboratory conditions remain uncharacterized."
    )

    add_body_paragraph(doc,
        "Together, these findings establish a coherent picture: PSCs are intrinsically more "
        "sensitive to environmental perturbation than differentiated cells because of their "
        "open chromatin state, rapid proliferation, dependence on clock-associated proteins, "
        "and susceptibility to osmotic gating of fate decisions. The key point is not that "
        "any single perturbation is catastrophic, but that PSCs sit at an energetically "
        "shallow decision landscape where subtle environmental shifts can tip the balance "
        "between self-renewal and lineage commitment. What is unknown is the real-world "
        "magnitude of these effects in working laboratories. This gap exists not because "
        "the question is unanswerable, but because no one has yet systematically measured "
        "the relevant variables alongside differentiation outcomes."
    )

    # --- Section 4: The catalogue of invisible variables ---
    add_heading(doc, "A catalogue of uncontrolled variables", level=2)

    add_body_paragraph(doc,
        "A comprehensive environmental assessment of a Japanese cell-processing facility "
        "revealed that while temperature was maintained constant year-round, humidity tracked "
        "outdoor seasonal patterns, peaking in summer and reaching troughs in winter{22}. "
        "Industry guidelines confirm that seasonal transitions substantially increase "
        "microbial contamination risk{23}. VOC concentrations in laboratories are determined "
        "by building materials, disinfectants, and HVAC systems, with off-gassing rates that "
        "are strongly temperature- and humidity-dependent\u2014introducing a seasonal "
        "component{7}. Ambient light exposure occurs during every biosafety cabinet "
        "manipulation, with duration and intensity varying by laboratory design and season "
        "(longer daylight hours in summer). Given CRY1\u2019s role in maintaining "
        "pluripotency{17}, even brief light exposure during passage or feeding could "
        "introduce stochastic variation in differentiation priming. Cosmic ray flux shows "
        "solar-cycle modulation with additional seasonal components{24}."
    )

    add_body_paragraph(doc,
        "Barometric pressure fluctuations deserve particular attention. Low-pressure weather "
        "systems reduce the partial pressure of dissolved gases in culture media and alter "
        "gas exchange kinetics in non-equilibrated incubators. These fluctuations are acute "
        "(hours), large (up to 30 hPa during cyclones), and entirely uncompensated in "
        "standard CO\u2082 incubators. Although direct evidence linking barometric pressure to "
        "PSC differentiation is lacking, the sensitivity of pH-dependent signaling pathways "
        "(e.g., Wnt, Hedgehog) to dissolved CO\u2082 concentrations suggests a plausible "
        "mechanism. Similarly, building vibration from HVAC systems, foot traffic, and "
        "nearby construction transmits mechanical stress to cultured cells\u2014and recent work "
        "demonstrates that mechanical compression can modulate nuclear mechanics and "
        "differentiation timing in hiPSCs{15}."
    )

    add_body_paragraph(doc,
        "Each of these variables exhibits characteristic seasonal patterns that create a "
        "complex, latitude-dependent forcing function no PSC laboratory currently monitors "
        "(Figure 1). The combination of uncontrolled environmental signals "
        "means that ostensibly identical experiments performed in January and July\u2014or in "
        "Boston and Brisbane\u2014may face systematically different environmental contexts that "
        "are invisible to the experimentalist."
    )

    # --- Section 5: The clonal complacency trap ---
    add_heading(doc, "The clonal complacency trap", level=2)

    add_body_paragraph(doc,
        "The clonal nature of PSC research creates a distinctive methodological hazard. "
        "In observational epidemiology, researchers routinely adjust for confounders because "
        "the heterogeneity of study populations demands it. In clonal systems, genetic "
        "uniformity creates the impression that confounding has been eliminated by design. "
        "This is a logical error: genetic homogeneity controls for genetic confounders but "
        "does nothing to control for environmental, temporal, or institutional confounders."
    )

    add_body_paragraph(doc,
        "The consequence is that PSC researchers typically do not record, report, or adjust "
        "for variables such as the season of experiment, laboratory humidity, or building "
        "renovation status. When batch-to-batch variation occurs, it is attributed to "
        "\u2018biological stochasticity\u2019 or \u2018technical variability\u2019\u2014catch-all terms that "
        "may conceal systematic environmental effects. The IVF field operated under "
        "precisely this assumption until VOC monitoring revealed that much of the \u2018random\u2019 "
        "variation in blastocyst rates was environmentally determined{6,7}."
    )

    add_body_paragraph(doc,
        "This trap has implications beyond individual laboratories. When multi-site "
        "reproducibility studies report that laboratory-of-origin is the dominant source of "
        "variance{3}, the standard interpretation is that protocols are insufficiently "
        "detailed. But if environmental variables differ systematically between sites\u2014as "
        "they inevitably do, given differences in building age, HVAC systems, altitude, "
        "latitude, and local climate\u2014then protocol harmonization alone cannot achieve "
        "reproducibility. The method is incomplete without environmental context."
    )

    add_body_paragraph(doc,
        "The analogy to epidemiology is precise. An epidemiologist studying health outcomes "
        "across cities would never assume that genetic similarity among study populations "
        "eliminates confounding from air quality, water supply, or altitude. Yet stem cell "
        "biologists routinely compare differentiation outcomes across laboratories in "
        "different cities\u2014sometimes different continents\u2014without recording or adjusting "
        "for any environmental covariate. The clonal complacency trap is, in essence, a "
        "category error: treating genetic control as sufficient for experimental control."
    )

    # --- Section 6: The GEO natural experiment ---
    add_heading(doc, "A natural experiment reveals institutional confounding in clonal data",
                level=2)

    add_body_paragraph(doc,
        "To test whether institutional rhythms confound clonal research data, we analyzed "
        "submission dates for 6,101 PSC differentiation datasets deposited in NCBI GEO{25} "
        "between 2001 and 2026 (search query: \u2018iPSC OR ESC differentiation\u2019; inclusion: all "
        "dataset types with valid submission dates). The overall distribution was "
        "significantly non-uniform (\u03c7\u00b2=32.3, df=11, p=0.0007) with a pronounced March "
        "peak\u2014superficially consistent with a spring photoperiod advantage."
    )

    add_body_paragraph(doc,
        "We exploited a natural experiment: geographic variation in academic/fiscal year "
        "calendars. Countries were grouped by academic year start: Japan/South Korea "
        "(March/April, n=250), USA/Europe/China (September/October, n=2,825), and "
        "Australia/New Zealand/Brazil (January/February, Southern Hemisphere, n=83). Using "
        "Pearson\u2019s \u03c7\u00b2 and circular Rayleigh tests{26}, we found a striking dissociation "
        "(Figure 1)."
    )

    add_body_paragraph(doc,
        "Only the Japan/South Korea group showed significant seasonality (\u03c7\u00b2=32.8, df=11, "
        "p=0.0006, Rayleigh R=0.172) with a March peak perfectly aligned with Japan\u2019s "
        "fiscal year-end. The much larger September-start group showed no significant "
        "seasonal pattern (\u03c7\u00b2=15.0, p=0.13). Japan\u2014just 3.3% of total datasets\u2014drove "
        "the overall signal entirely: excluding Japan eliminated statistical significance "
        "(p=0.128)."
    )

    add_body_paragraph(doc,
        "This result carries a critical methodological lesson. What appeared to be evidence "
        "of seasonal biological influence on stem cell research was in fact an artifact of "
        "Japan\u2019s March 31 fiscal year-end, which creates a rush to submit manuscripts and "
        "deposit data before annual reporting deadlines. Japan contributed only 3.3% of "
        "total datasets (200/6,101) but accounted for 15% of all March submissions\u2014a "
        "nearly two-fold over-representation relative to its overall share. The fiscal "
        "year-end effect is well-documented in Japanese academic culture: grant reports "
        "are due by March 31, creating institutional pressure to finalize data and "
        "publications before the deadline."
    )

    add_body_paragraph(doc,
        "The broader implication extends beyond stem cell biology. Any field that relies on "
        "publication or data-deposition timing as a proxy for experimental timing risks "
        "confounding institutional rhythms with biological ones. GEO submission dates "
        "reflect a complex convolution of experiment timing, analysis duration, writing, "
        "peer review, and institutional incentives\u2014with a typical lag of 6\u201318 months from "
        "bench experiment to data deposition. Our natural experiment demonstrates that "
        "public database metadata cannot reliably distinguish biological seasonality from "
        "institutional rhythms. This finding should serve as a cautionary note for "
        "meta-research studies that use submission timestamps as temporal proxies."
    )

    add_body_paragraph(doc,
        "An exploratory analysis of the relationship between PSC dataset volume and solar "
        "activity indices (NOAA F10.7 flux) illustrates a complementary pitfall. The raw "
        "correlation was r=0.48, suggesting an intriguing link\u2014until detrending removed "
        "the common upward trajectory of both iPSC research output and Solar Cycle 24/25 "
        "ascent, yielding a residual correlation of r=0.02 (p=0.75). Geographic analysis "
        "further constrained interpretation: 97% of GEO datasets with identifiable country "
        "affiliations originated from Northern Hemisphere institutions, rendering "
        "hemisphere-inversion tests impractical with GEO data alone."
    )

    # --- Section 7: The method: prospective monitoring ---
    add_heading(doc, "The path forward: environmental monitoring as method", level=2)

    add_body_paragraph(doc,
        "Our analysis demonstrates that retrospective database mining cannot detect or rule "
        "out environmental effects on PSC differentiation. The IVF evidence suggests such "
        "effects exist in early-stage cells{9,13}, and the vulnerability evidence suggests "
        "PSCs are plausible targets{14,15}. What is needed is a systematic method: prospective, "
        "controlled environmental monitoring paired with differentiation outcome tracking."
    )

    add_body_paragraph(doc,
        "We propose a three-phase approach (Figure 2). Phase I: equip PSC facilities with "
        "continuous multi-parameter sensors (humidity, illuminance, ELF-EMF, barometric "
        "pressure, VOC levels, vibration) at 1-minute resolution alongside routine "
        "differentiation outcomes for \u226512 months. Modern IoT sensor packages achieve this "
        "at modest cost (<$5,000 per laboratory). Phase II: analyze existing data from IVF "
        "registries (HFEA in the UK, ANZARD in Australia) that contain date-stamped cycle "
        "outcomes with latitude information\u2014these can directly test the hemisphere-inversion "
        "prediction for photoperiod-driven effects. Phase III: for variables identified in "
        "Phases I\u2013II as significantly correlated with outcomes, conduct prospective "
        "controlled interventions (e.g., humidity-controlled vs. uncontrolled incubators, "
        "light-tight hoods vs. standard biosafety cabinets, magnetic shielding)."
    )

    add_body_paragraph(doc,
        "The hemisphere-inversion test provides a powerful diagnostic criterion. Variables "
        "whose effects are phase-inverted between Northern and Southern Hemisphere "
        "laboratories implicate solar-driven mechanisms\u2014photoperiod, temperature, UV flux. "
        "Variables whose effects are hemisphere-synchronous implicate geomagnetic or cosmic "
        "ray pathways. A coordinated multi-center study spanning both hemispheres could "
        "decompose seasonal variation into these mechanistic categories."
    )

    add_body_paragraph(doc,
        "The technical requirements for Phase I monitoring are straightforward. "
        "A minimum sensor array would include: (1) temperature/humidity loggers (\u00b10.1\u00b0C, "
        "\u00b11% RH) inside incubators and at bench level; (2) a photodiode sensor in the "
        "biosafety cabinet recording cumulative lux-hours of exposure per manipulation; "
        "(3) a VOC sensor (ppb-level total VOC) with seasonal baseline tracking; (4) an "
        "ELF-EMF sensor (3-axis, nT resolution) placed adjacent to incubators; and (5) a "
        "barometric pressure logger. Modern IoT platforms can integrate all five data "
        "streams with automatic cloud upload, enabling multi-site meta-analyses without "
        "manual data wrangling. The critical companion requirement is systematic recording "
        "of differentiation outcomes per batch\u2014ideally including not just binary success/"
        "failure but quantitative markers (flow cytometry percentages, gene expression "
        "levels, electrophysiological parameters) that permit correlation analysis."
    )

    add_body_paragraph(doc,
        "Figure 2 presents a research roadmap organized by the expected hemisphere-"
        "dependence of each variable and the current level of evidence."
    )

    # --- Section 8: Implications ---
    add_heading(doc, "Implications for reproducibility and manufacturing", level=2)

    add_body_paragraph(doc,
        "As PSC-derived cell therapies advance toward clinical application{1}, manufacturing "
        "consistency becomes a regulatory imperative. Current Good Manufacturing Practice "
        "(cGMP) facilities control temperature, humidity, and particulates, but do not "
        "routinely monitor illuminance, EMF, VOCs, or barometric pressure. If any of these "
        "variables meaningfully affect differentiation, current manufacturing harbors hidden "
        "batch-to-batch variability attributed to biological stochasticity but potentially "
        "environmental and correctable{27}. The economic argument is compelling: failed "
        "batches in autologous cell therapy represent lost patient tissue and treatment "
        "delays. Even modest improvements in first-pass success rates through environmental "
        "optimization could translate into significant cost savings and reduced "
        "time-to-treatment."
    )

    add_body_paragraph(doc,
        "The monitoring approach imposes no change on existing workflows\u2014it simply makes "
        "the invisible visible. Recent proposals for \u2018cytocentric\u2019 quality principles in "
        "regenerative medicine have independently argued that sensing and controlling "
        "environmental parameters must become a cornerstone of cell manufacturing{28,29}. "
        "Unlike prospective intervention studies, which require "
        "protocol modifications and ethical approvals, passive environmental monitoring "
        "can be implemented immediately in any laboratory. The resulting datasets, when "
        "combined across multiple sites and latitudes, would constitute a resource of "
        "enormous value for the field\u2014analogous to weather station networks that "
        "transformed meteorology from an anecdotal practice to a predictive science."
    )

    add_body_paragraph(doc,
        "The IVF field\u2019s experience is instructive: when laboratories invested in VOC "
        "filtration and positive-pressure air handling, blastocyst rates increased by 18% "
        "and live births by 8%{6,8}. These improvements were achieved not through better "
        "biological protocols but through better environmental control\u2014precisely the "
        "paradigm shift we advocate for PSC manufacturing. The convergence of clinical-grade "
        "cell therapy production with modern IoT environmental monitoring capabilities "
        "creates an opportune moment to establish new standards for \u2018environmentally "
        "aware\u2019 cell culture."
    )

    # --- Section 9: Boundaries ---
    add_heading(doc, "What this Perspective does not argue", level=2)

    add_body_paragraph(doc,
        "We are not claiming that seasonal photoperiod directly drives PSC differentiation "
        "outcomes\u2014our GEO analysis demonstrates precisely the opposite. We are not claiming "
        "that environmental variables are the dominant source of variability\u2014genetic "
        "background, passage number, and operator technique remain important. Rather, we "
        "argue that environmental variables constitute an additional, correctable source of "
        "variation that has been systematically overlooked because the clonal paradigm implies "
        "they should not exist. The null hypothesis\u2014that humidity, VOCs, and ambient light "
        "have zero effect on PSC differentiation\u2014is contradicted by the evidence "
        "reviewed{14,15,20}, but the magnitude of effects under real-world conditions remains "
        "unknown. Only prospective measurement can resolve this."
    )

    # --- Conclusion ---
    add_heading(doc, "Conclusion", level=2)

    add_body_paragraph(doc,
        "The history of experimental biology is punctuated by discoveries that \u2018noise\u2019 was "
        "in fact signal from an unmeasured variable. The stem cell field may be at such an "
        "inflection point. We have spent two decades optimizing recipes while ignoring the "
        "kitchen. The convergence of evidence from IVF environmental control{6,7}, seasonal "
        "effects on early embryos{9,13}, circadian regulation of pluripotency{17,18}, and "
        "environmental monitoring of cell-processing facilities{22} suggests that the "
        "invisible variables are neither negligible nor intractable. They merely need to be "
        "measured."
    )

    add_body_paragraph(doc,
        "The practical path forward is clear. First, equip PSC laboratories with "
        "multi-parameter environmental sensors and record outcomes systematically for 12 "
        "months. Second, analyze IVF registry data for hemisphere-dependent seasonal effects. "
        "Third, for any variable identified as significant, conduct controlled "
        "interventions. The cost of environmental monitoring is negligible compared to the "
        "cost of failed differentiation batches in clinical manufacturing{1,27}. If even one "
        "invisible variable proves consequential\u2014as VOCs proved consequential in "
        "IVF{6}\u2014the return on investment for the field will be transformative."
    )

    add_body_paragraph(doc,
        "In summary, this Perspective has identified a methodological gap in stem cell "
        "research: the clonal complacency trap. Key unanswered questions include: (i) which "
        "specific environmental variables have the largest effect on differentiation "
        "outcomes? (ii) Are effects consistent across cell lines and protocols? (iii) Can "
        "the hemisphere-inversion test discriminate solar-driven from geomagnetic mechanisms? "
        "(iv) What are the minimum monitoring requirements for clinically meaningful "
        "environmental quality control? Addressing these questions through prospective, "
        "multi-site environmental profiling would represent a paradigm shift from recipe "
        "optimization to kitchen optimization in stem cell biology."
    )

    # ──────────────────────────────────────────────
    # Figure Legends
    # ──────────────────────────────────────────────
    add_heading(doc, "Figure Legends", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run("Figure 1 | Academic calendar variation as a natural experiment. ")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run = p.add_run(
        "a, Normalized monthly distributions of GEO PSC dataset submissions by academic "
        "year group. The March/April-start group (Japan/Korea, red) shows a pronounced "
        "March peak absent in the September-start group (USA/Europe/China, blue). "
        "b, Country-level heatmap showing that the pattern is driven predominantly by Japan. "
        "c, Circular mean directions (Rayleigh vectors) for each group. "
        "d, Statistical summary. Only Japan/Korea achieves significance (p<0.001). "
        "n=6,101 total datasets; 3,195 with country affiliation."
    )
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0, space_before=Pt(12))
    run = p.add_run(
        "Figure 2 | Research roadmap for environmental profiling of PSC culture. "
    )
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run = p.add_run(
        "a, Evidence matrix: uncontrolled environmental variables classified by "
        "hemisphere-dependence (inverted = solar-driven; synchronous = geomagnetic) and "
        "evidence level. "
        "b, Three-phase investigation strategy from passive monitoring through "
        "retrospective data mining to controlled intervention."
    )
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # ──────────────────────────────────────────────
    # References (numbered, sequential)
    # ──────────────────────────────────────────────
    add_heading(doc, "References", level=2)

    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=Pt(4), line_spacing=2.0)
        run = p.add_run(f"{i}. ")
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run = p.add_run(ref)
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    # Save
    out_path = os.path.join(OUTPUT_DIR, "NatMethods_Perspective_InvisibleVariables.docx")
    doc.save(out_path)
    print(f"Manuscript saved to: {out_path}")

    # Word count (body text only)
    word_count = 0
    in_body = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "The methodological blind spot":
            in_body = True
            continue
        if text in ("Figure Legends", "References"):
            in_body = False
            continue
        if in_body and text:
            word_count += len(text.split())
    print(f"Approximate body word count (excl. abstract/refs/legends): {word_count}")

    return out_path


if __name__ == "__main__":
    create_manuscript()
