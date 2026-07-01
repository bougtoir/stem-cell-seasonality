#!/usr/bin/env python3
"""
Generate cover letter for Human Reproduction Opinion submission.

Requirements (per HR guidelines):
- Addressed to the Editor-in-Chief
- State the scientific question asked
- State the principal new findings
- State their significance to the field
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_cover_letter():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_para(text, space_after=Pt(6), bold=False, italic=False,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = alignment
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
        return p

    # Date
    today = date.today()
    add_para(today.strftime("%B %d, %Y"), space_after=Pt(18))

    # Addressee
    add_para("Editor-in-Chief", space_after=Pt(0))
    add_para("Human Reproduction", space_after=Pt(0), italic=True)
    add_para("European Society of Human Reproduction and Embryology (ESHRE)",
             space_after=Pt(18))

    # Salutation
    add_para("Dear Editor-in-Chief,", space_after=Pt(12))

    # Body
    add_para(
        "We respectfully submit our manuscript entitled "
        "\"The Invisible Variables: Uncontrolled Environmental Factors in Stem Cell "
        "Differentiation and the Fiscal-Year Artifact in Public Databases\" "
        "for consideration as an Opinion article in Human Reproduction.",
        space_after=Pt(12)
    )

    add_para(
        "Scientific question: Can uncontrolled environmental variables in stem cell "
        "laboratories\u2014humidity, volatile organic compounds (VOCs), ambient light, "
        "electromagnetic fields, and barometric pressure\u2014explain a meaningful portion of "
        "the batch-to-batch variability that plagues pluripotent stem cell (PSC) "
        "differentiation, and does the apparent seasonal pattern in public research "
        "databases reflect a biological signal or an institutional artifact?",
        space_after=Pt(12)
    )

    add_para(
        "Principal new findings: We present three convergent lines of evidence. First, "
        "we synthesize recent literature demonstrating that PSCs are uniquely vulnerable to "
        "environmental perturbations at magnitudes that leave somatic cells unaffected\u2014"
        "including osmotic stress, airborne pollutants, and circadian disruption. Second, "
        "we analyze 6,101 GEO PSC datasets and use a natural experiment exploiting geographic "
        "variation in academic year calendars (Japan March/April start vs. USA/Europe "
        "September start vs. Southern Hemisphere January start) to demonstrate that the "
        "apparent March peak in dataset submissions is driven entirely by Japan's fiscal "
        "year-end\u2014an institutional artifact, not a biological signal. Third, we propose a "
        "concrete three-phase research roadmap for prospective environmental monitoring "
        "of PSC culture facilities.",
        space_after=Pt(12)
    )

    add_para(
        "Significance to the field: This work is directly relevant to Human Reproduction's "
        "readership for several reasons. (1) Our argument draws extensively on the IVF "
        "field's experience: Leathersich et al. (2023, Hum Reprod) showed that oocytes "
        "collected in summer yield 30% higher live birth rates, and Cai et al. (2025, "
        "Hum Reprod) demonstrated that VOC reduction improved blastocyst rates by 18%. "
        "We use these findings as proof-of-concept that environmental variables affect "
        "early-stage cell biology. (2) The methodological lesson\u2014that public database "
        "\"seasonality\" can be an institutional artifact\u2014is directly applicable to IVF "
        "registry studies that use treatment timing as a variable. (3) Our proposed "
        "monitoring framework, including hemisphere-inversion tests using HFEA and ANZARD "
        "registry data, represents actionable next steps for the reproductive medicine "
        "community.",
        space_after=Pt(12)
    )

    add_para(
        "This manuscript has not been published previously, is not under consideration "
        "elsewhere, and all authors have approved the submitted version. We have no "
        "conflicts of interest to declare.",
        space_after=Pt(12)
    )

    add_para(
        "We believe this Opinion will stimulate important discussion about environmental "
        "quality control in both stem cell and IVF laboratories, and we hope you will find "
        "it suitable for Human Reproduction.",
        space_after=Pt(12)
    )

    # Closing
    add_para("Sincerely,", space_after=Pt(24))
    add_para("[Corresponding author name]", space_after=Pt(0))
    add_para("[Affiliation]", space_after=Pt(0))
    add_para("[Email address]", space_after=Pt(0))
    add_para("[ORCID iD]", space_after=Pt(0))

    # Save
    out_path = os.path.join(OUTPUT_DIR, "HumReprod_CoverLetter.docx")
    doc.save(out_path)
    print(f"Cover letter saved to: {out_path}")

    return out_path


if __name__ == "__main__":
    create_cover_letter()
