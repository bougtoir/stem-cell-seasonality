#!/usr/bin/env python3
"""
Generate presubmission inquiry letter for Nature Methods.

Nature Methods presubmission inquiry process:
- Submit via Manuscript Tracking System: https://mts-nmeth.nature.com/
- Provide abstract + brief rationale for why the work fits Nature Methods
- Do NOT attach full manuscript at inquiry stage
- Editorial contact: methods@us.nature.com
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_inquiry():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("[Date]")
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # Addressee
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Dr. Allison Doerr\nChief Editor\nNature Methods")
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # Salutation
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Dear Dr. Doerr,")
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # Body paragraphs
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "We write to inquire whether Nature Methods would consider a Perspective article "
        "entitled \u201cThe Invisible Variables: Why Clonal Systems Are Not Immune to "
        "Environmental Confounding.\u201d"
    )
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "Pluripotent stem cell (PSC) differentiation protocols suffer persistent "
        "batch-to-batch variability that limits both basic research reproducibility and "
        "clinical manufacturing consistency. We identify a methodological blind spot: "
        "multiple environmental variables\u2014humidity, volatile organic compounds, ambient "
        "light, electromagnetic fields, and barometric pressure\u2014remain unmonitored in "
        "most stem cell laboratories. We introduce the concept of the \u2018clonal complacency "
        "trap\u2019: the widespread assumption that genetic homogeneity in clonal systems "
        "eliminates the need for environmental confounding adjustment."
    )
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "We draw on two precedents where environmental control transformed outcomes: "
        "artisanal baking and IVF embryology (where VOC reduction increased blastocyst "
        "rates by 18%). A natural experiment using 6,101 GEO datasets demonstrates that "
        "apparent seasonality in public stem cell databases is an institutional artifact "
        "(driven by Japan\u2019s March 31 fiscal year-end), not a biological signal\u2014providing "
        "a cautionary lesson for all database-mining meta-research. We then review evidence "
        "that PSCs are uniquely vulnerable to environmental perturbation (drug sensitivity "
        "gradients, mechano-osmotic gating of fate decisions, non-canonical circadian "
        "clock dependence) and propose IoT-based multi-parameter environmental monitoring "
        "as a systematic method to identify which invisible variables matter."
    )
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "We believe this work is well-suited for Nature Methods because it identifies a "
        "fundamental methodological gap affecting reproducibility across the stem cell "
        "field and proposes a concrete, actionable monitoring framework. The \u2018clonal "
        "complacency trap\u2019 concept is broadly relevant to any field using isogenic "
        "systems, and the GEO fiscal-year artifact has implications for database-mining "
        "studies across biology. The manuscript is approximately 3,300 words with 26 "
        "references and 2 display items."
    )
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(
        "This manuscript has not been published previously and is not under consideration "
        "elsewhere. We would be pleased to provide the full manuscript or any additional "
        "information."
    )
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    # Closing
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Sincerely,")
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run(
        "[Corresponding author name]\n"
        "[Affiliation]\n"
        "[Email]\n"
        "[ORCID iD]"
    )
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.italic = True

    # Save
    out_path = os.path.join(OUTPUT_DIR,
                             "NatMethods_Presubmission_Inquiry.docx")
    doc.save(out_path)
    print(f"Presubmission inquiry saved to: {out_path}")
    return out_path


if __name__ == "__main__":
    create_inquiry()
