#!/usr/bin/env python3
"""
Generate cover letter for STEM CELLS Perspective submission.

Requirements (per STEM CELLS guidelines):
- Briefly describe the work's significance
- Identify corresponding author with:
  - Complete mailing address
  - Telephone number
  - Email address
- Disclose related preprints/manuscripts/papers
"""

import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def create_cover_letter():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_para(text, space_after=Pt(6), bold=False, italic=False,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = alignment
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
        return p

    # Date
    today = date.today()
    add_para(today.strftime("%B %d, %Y"), space_after=Pt(18))

    # Addressee
    add_para("Editor-in-Chief", space_after=Pt(0))
    add_para("Stem Cells", space_after=Pt(0), italic=True)
    add_para("Oxford University Press", space_after=Pt(18))

    # Salutation
    add_para("Dear Editor,", space_after=Pt(12))

    # Opening paragraph
    add_para(
        "We respectfully submit our manuscript entitled "
        "\u201cThe Invisible Variables: Why Clonal Systems Are Not Immune to "
        "Environmental Confounding\u201d "
        "for consideration as a Perspective in Stem Cells.",
        space_after=Pt(12)
    )

    # Significance
    add_para(
        "Pluripotent stem cell (PSC)-derived therapies are now entering clinical "
        "trials at scale, yet batch-to-batch variability in differentiation "
        "protocols remains a major barrier to both research reproducibility and "
        "clinical manufacturing. Our Perspective identifies an overlooked source "
        "of this variability: unmonitored environmental variables in the "
        "laboratory\u2014including humidity, volatile organic compounds (VOCs), "
        "ambient light, electromagnetic fields, and barometric pressure. We "
        "introduce the concept of the \u2018clonal complacency trap\u2019: the assumption "
        "that genetic uniformity in isogenic cell lines guarantees experimental "
        "uniformity, which discourages investigators from measuring or adjusting "
        "for non-genetic confounders.",
        space_after=Pt(12)
    )

    add_para(
        "We present three convergent lines of evidence. First, we draw on the "
        "IVF field\u2019s transformative experience, where controlling air quality and "
        "VOC exposure improved blastocyst rates by up to 18%, demonstrating that "
        "environmental optimization of early-stage cell biology is both feasible "
        "and impactful. Second, we synthesize recent findings showing that PSCs "
        "are uniquely vulnerable to environmental perturbations\u2014osmotic stress, "
        "circadian disruption, and airborne pollutants\u2014at magnitudes that leave "
        "somatic cells unaffected. Third, we analyze 6,101 GEO PSC datasets in "
        "a natural experiment exploiting geographic variation in academic "
        "calendars, demonstrating that apparent \u2018seasonality\u2019 in public databases "
        "is an institutional artifact driven by Japan\u2019s fiscal year-end, not a "
        "biological signal. This underscores the need for prospective "
        "environmental measurement rather than retrospective inference.",
        space_after=Pt(12)
    )

    add_para(
        "We propose a concrete, low-cost research roadmap centered on IoT-based "
        "multi-parameter environmental monitoring to identify which invisible "
        "variables matter for PSC culture and cGMP manufacturing. This framework "
        "is directly actionable for the stem cell community.",
        space_after=Pt(12)
    )

    # Disclosure of related work
    add_para(
        "Disclosure of related work: We have no related preprints, manuscripts, "
        "or published papers containing substantially similar content or using "
        "the same data currently under consideration elsewhere.",
        space_after=Pt(12)
    )

    # Ethics and COI
    add_para(
        "This manuscript has not been published previously and is not under "
        "consideration elsewhere. All authors have approved the submitted "
        "version and have no conflicts of interest to declare. No human or "
        "animal subjects were involved; the GEO analysis used only publicly "
        "available metadata.",
        space_after=Pt(12)
    )

    add_para(
        "We believe this Perspective will stimulate important discussion about "
        "environmental quality control in stem cell laboratories and the "
        "implications for clinical-grade manufacturing. We hope you will find "
        "it suitable for Stem Cells.",
        space_after=Pt(12)
    )

    # Closing
    add_para("Sincerely,", space_after=Pt(24))

    # Corresponding author block
    add_para("[Corresponding author name]", space_after=Pt(0), bold=True)
    add_para("[Affiliation / Department]", space_after=Pt(0))
    add_para("[Complete mailing address]", space_after=Pt(0))
    add_para("[Telephone number]", space_after=Pt(0))
    add_para("[Email address]", space_after=Pt(0))
    add_para("[ORCID iD]", space_after=Pt(0))

    # Save
    out_path = os.path.join(OUTPUT_DIR, "StemCells_CoverLetter.docx")
    doc.save(out_path)
    print(f"Cover letter saved to: {out_path}")

    return out_path


if __name__ == "__main__":
    create_cover_letter()
