#!/usr/bin/env python3
"""
Generate Stem Cells Perspective manuscript:
"The Invisible Variables: Why Clonal Systems Are Not Immune to
 Environmental Confounding"

Format: STEM CELLS Perspective (Oxford Academic / Wiley)
- Word limit: 2,000 words (excluding abstract, tables, figures, legends, references)
- Display items: 2 (figures/illustrations/tables)
- Max references: 50
- Abstract: unstructured, <= 250 words
- Significance Statement: required (~10 sentences, lay-reader friendly)
- Reference style: AMA (numbered in order of appearance)
- US spelling throughout
- Sections: flexible for Perspectives (not strict IMRaD)

Restructured from Nature Methods Perspective (4,000 words) to fit
Stem Cells 2,000-word limit while preserving core argument.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ──────────────────────────────────────────────
# References in order of first appearance in text
# AMA style: Author AB, Author CD, Author EF.
#   Title. J Abbrev. Year;Vol(Issue):Pages. doi:
# ──────────────────────────────────────────────
REFERENCES = [
    # 1 - Kirkeby (clinical trials update) [Introduction]
    "Kirkeby A, Main H, Carpenter M. Pluripotent stem-cell-derived "
    "therapies in clinical trial: a 2025 update. Cell Stem Cell. 2025;32(1):10-37.",
    # 2 - Yamanaka (PSC therapy challenges) [Introduction]
    "Yamanaka S. Pluripotent stem cell-based cell therapy\u2014promise and challenges. "
    "Cell Stem Cell. 2020;27(4):523-531.",
    # 3 - Volpato (multi-site reproducibility) [Introduction]
    "Volpato V, Smith J, Sandor C, et al. Reproducibility of molecular phenotypes "
    "after long-term differentiation to human iPSC-derived neurons: a multi-site omics "
    "study. Stem Cell Reports. 2018;11(4):897-911.",
    # 4 - Volpato & Webber (iPSC variability guidelines) [Introduction]
    "Volpato V, Webber C. Addressing variability in iPSC-derived models of human "
    "disease: guidelines to promote reproducibility. Dis Model Mech. 2020;13(1):dmm042317.",
    # 5 - Mortimer (Cairo consensus IVF lab air quality) [IVF precedent]
    "Mortimer D, Cohen J, Mortimer ST, et al. Cairo consensus on the IVF laboratory "
    "environment and air quality: report of an expert meeting. Reprod Biomed Online. "
    "2018;36(6):658-674.",
    # 6 - Agarwal (VOCs and IVF outcomes) [IVF precedent]
    "Agarwal N, Chattopadhyay R, Chakravarty B. Volatile organic compounds and good "
    "laboratory practices in the in vitro fertilization laboratory: the important "
    "parameters for successful outcome in extended culture. "
    "J Assist Reprod Genet. 2017;34(8):999-1006.",
    # 7 - Cai (VOC real-time monitoring, embryology lab) [IVF precedent]
    "Cai J, Zhou H, Wang M, et al. Real-time monitoring reveals the effects of low "
    "concentrations of volatile organic compounds in the embryology laboratory. "
    "Hum Reprod. 2025;40(3):deaf008.",
    # 8 - Leathersich (IVF season and oocyte collection) [IVF seasonality]
    "Leathersich SJ, Roche CS, Walls M, Nathan E, Hart RJ. Season at the time "
    "of oocyte collection and frozen embryo transfer outcomes. Hum Reprod. "
    "2023;38(9):1714-1722.",
    # 9 - Wang C (IVF season, temperature, DNI) [IVF seasonality]
    "Wang C, Chen J, Lin Z, et al. The impact of season, temperature, and direct "
    "normal irradiance on IVF pregnancy outcomes: a retrospective cohort study. "
    "Int J Biometeorol. 2025;69(10):2053-2065.",
    # 10 - Wang JP (mouse seasonal two-cell block) [IVF seasonality]
    "Wang JP, Her WY, Meir YJ, Liu TS, Chang HL, Haung FL. Seasonal variation in "
    "cell cycle during early development of the mouse embryo. Reproduction. "
    "1992;94(2):431-436.",
    # 11 - Panina (PSC drug sensitivity) [PSC vulnerability]
    "Panina Y, Yamane J, Kobayashi K, Sone H, Fujibuchi W. Human ES and iPS cells "
    "display less drug resistance than differentiated cells, and naive-state induction "
    "further decreases drug resistance. J Toxicol Sci. 2021;46(3):131-142.",
    # 12 - McCreery (mechano-osmotic chromatin) [PSC vulnerability]
    "McCreery KP, Stubb A, Stephens R, et al. Mechano-osmotic signals control "
    "chromatin state and fate transitions in pluripotent stem cells. Nat Cell Biol. "
    "2025;27(10):1757-1770.",
    # 13 - Chui (osmolar modulation, iPSC differentiation) [PSC vulnerability]
    "Chui JSH, Izuel-Idoype T, Qualizza A, et al. Osmolar modulation drives "
    "reversible cell cycle exit and human pluripotent cell differentiation via "
    "NF-\u03baB and WNT signaling. Adv Sci. 2024;11(7):2307554.",
    # 14 - Sato (CRY1 in PSCs) [PSC vulnerability]
    "Sato S, Hishida T, Kinouchi K, et al. The circadian clock CRY1 regulates "
    "pluripotent stem cell identity and somatic cell reprogramming. Cell Rep. "
    "2023;42(6):112590.",
    # 15 - Ameneiro (BMAL1 in PSCs) [PSC vulnerability]
    "Ameneiro C, Moreira T, Fuentes-Iglesias A, et al. BMAL1 coordinates energy "
    "metabolism and differentiation of pluripotent stem cells. Life Sci Alliance. "
    "2020;3(5):e201900534.",
    # 16 - Bi (PM2.5 and hESC pluripotency) [PSC vulnerability]
    "Bi S, Liang W, Qi R, et al. Fine particulate matter reduces the pluripotency "
    "and proliferation of human embryonic stem cells through ROS induced AKT and ERK "
    "signaling pathway. Reprod Toxicol. 2020;96:231-240.",
    # 17 - Barrett (NCBI GEO database) [GEO analysis]
    "Barrett T, Wilhite SE, Ledoux P, et al. NCBI GEO: archive for functional "
    "genomics data sets\u2014update. Nucleic Acids Res. 2013;41(D1):D991-D995.",
    # 18 - Mardia (directional/circular statistics) [GEO analysis]
    "Mardia KV, Jupp PE. Directional Statistics. Wiley; 2000.",
    # 19 - Klein (best practices cell culture environments) [Monitoring]
    "Klein SG, Alsolami SM, Steckbauer A, et al. Toward best practices for "
    "controlling mammalian cell culture environments. Front Cell Dev Biol. "
    "2022;10:788808.",
    # 20 - Karagiannis (iPSC use in disease models) [Implications/Manufacturing]
    "Karagiannis P, Takahashi K, Saito M, et al. Induced pluripotent stem cells "
    "and their use in human models of disease and development. Physiol Rev. "
    "2019;99(1):79-114.",
    # 21 - Henn (cytocentric principles for RM) [Implications]
    "Henn AD, Mitra K, Hunsberger J, et al. Applying the cytocentric principles to "
    "regenerative medicine for reproducibility. Curr Stem Cell Rep. 2022;8(4):197-202.",
    # 22 - Henn (cytocentric measurement for RM) [Implications]
    "Henn AD, Pereira T, Hunsberger J, et al. Cytocentric measurement for "
    "regenerative medicine. Front Med Technol. 2023;5:1154653.",
]


def add_superscript_refs(paragraph, text):
    """Parse text with {N} or {N-M} or {N,M} markers and create superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
    return paragraph


def set_paragraph_format(para, space_after=Pt(6), space_before=Pt(0),
                         line_spacing=2.0, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    para.paragraph_format.line_spacing = line_spacing
    para.alignment = alignment
    return para


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading


def add_body_paragraph(doc, text):
    """Add a body paragraph with superscript citation handling."""
    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    add_superscript_refs(p, text)
    return p


def insert_figure(doc, image_path, caption_text, fig_number):
    """Insert a figure inline with caption below."""
    if not os.path.exists(image_path):
        p = doc.add_paragraph()
        run = p.add_run(f"[Figure {fig_number} placeholder - source not found]")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.italic = True
        return
    # Add spacing before figure
    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(18), space_after=Pt(6))
    run = p.add_run()
    run.add_picture(image_path, width=Inches(5.5))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Caption
    cap = doc.add_paragraph()
    set_paragraph_format(cap, space_before=Pt(12), space_after=Pt(12))
    run = cap.add_run(f"Figure {fig_number}. ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run = cap.add_run(caption_text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'


def create_manuscript():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ──────────────────────────────────────────────
    # Title page
    # ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("PERSPECTIVE")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.bold = True

    p = doc.add_paragraph()
    set_paragraph_format(p, space_before=Pt(24), space_after=Pt(12))
    run = p.add_run(
        "The Invisible Variables: Why Clonal Systems Are Not Immune "
        "to Environmental Confounding"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(12))
    run = p.add_run("[Author names]")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(6))
    run = p.add_run("[Affiliations]")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True

    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(6))
    run = p.add_run("Correspondence: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run("[Corresponding author email]")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.italic = True

    # Keywords
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(6))
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run(
        "pluripotent stem cells; cell culture environment; reproducibility; "
        "environmental monitoring; differentiation variability"
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Classification
    p = doc.add_paragraph()
    set_paragraph_format(p, space_after=Pt(24))
    run = p.add_run("Classification: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run("Embryonic Stem Cells/Induced Pluripotent Stem (iPS) Cells")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # ──────────────────────────────────────────────
    # Abstract (<=250 words)
    # ──────────────────────────────────────────────
    add_heading(doc, "Abstract", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run(
        "Pluripotent stem cell (PSC) differentiation protocols suffer from persistent "
        "batch-to-batch variability that impedes both basic research and clinical "
        "manufacturing. We argue that this variability has a methodological blind spot: "
        "environmental variables\u2014humidity, volatile organic compounds, ambient light, "
        "electromagnetic fields, and barometric pressure\u2014remain unmonitored in most "
        "laboratories. The clonal nature of PSC research creates a false sense of "
        "experimental control; genetic homogeneity does not eliminate environmental "
        "confounding. We term this the 'clonal complacency trap'\u2014the assumption that "
        "genetic uniformity guarantees experimental uniformity. "
        "Drawing on the in vitro fertilization field's experience, where volatile organic "
        "compound control increased blastocyst rates by 18%, we illustrate how environmental "
        "optimization transformed outcomes. A natural experiment using 6,101 Gene Expression Omnibus (GEO) datasets "
        "demonstrates that apparent seasonality in public databases is an institutional "
        "artifact driven by Japan's fiscal year-end, not a biological signal\u2014highlighting "
        "the need for prospective measurement rather than retrospective inference. "
        "We propose IoT-based multi-parameter environmental monitoring as a systematic, "
        "low-cost method to identify which invisible variables matter for PSC culture."
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # ──────────────────────────────────────────────
    # Significance Statement
    # ──────────────────────────────────────────────
    add_heading(doc, "Significance Statement", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run(
        "Stem cell therapies are entering clinical trials, but batch-to-batch variability "
        "in differentiation remains a major barrier. We identify an overlooked source: "
        "the laboratory environment. Air quality, humidity, light, and barometric pressure "
        "fluctuate seasonally and differ between buildings, yet go unmeasured in stem cell "
        "experiments. Clonal cell lines create a false sense of control, masking these "
        "environmental influences. Fertility medicine solved a similar problem by "
        "controlling air quality, dramatically improving outcomes. We propose continuous "
        "environmental monitoring to reveal correctable sources of variability, improving "
        "research reproducibility and cell therapy manufacturing."
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # ──────────────────────────────────────────────
    # Main text (target: ~2,000 words)
    # ──────────────────────────────────────────────

    # --- Introduction ---
    add_heading(doc, "Introduction", level=2)

    add_body_paragraph(doc,
        "Pluripotent stem cell (PSC)-derived therapies are advancing rapidly, with over "
        "115 trials and 1,200 patients dosed as of 2024.{1} Yet differentiation protocols "
        "remain notoriously variable across laboratories and even between batches within a "
        "single group.{2-4} Multi-site reproducibility studies consistently report that "
        "laboratory-of-origin is the dominant source of variance in iPSC-derived neuron "
        "transcriptomes, exceeding the effect of genetic background or differentiation "
        "batch.{3} The field has invested enormously in optimizing controllable recipe "
        "variables\u2014growth factors, small molecules, extracellular matrices, timing\u2014while "
        "largely ignoring the environment in which these recipes are executed."
    )

    add_body_paragraph(doc,
        "A deeper methodological issue compounds the problem. Research using clonal "
        "systems\u2014isogenic iPSC lines, genetically identical clones\u2014operates under the "
        "implicit assumption that genetic homogeneity renders confounding adjustment "
        "unnecessary. Because the genetic background is 'controlled,' investigators rarely "
        "apply the stratification or sensitivity analyses that epidemiologists consider "
        "mandatory for observational data. Non-genetic confounders (environmental, "
        "temporal, institutional) are neither measured nor adjusted for. We term this the "
        "'clonal complacency trap'\u2014the assumption that genetic uniformity guarantees "
        "experimental uniformity."
    )

    add_body_paragraph(doc,
        "Consider an analogy. Artisanal bread quality varied unpredictably between batches "
        "until industrial bakeries recognized that controlling oven temperature alone was "
        "insufficient; ambient humidity, dough temperature during proofing, and air quality "
        "were equally critical. We argue that today's PSC laboratory resembles the "
        "pre-industrial bakery: we control incubator temperature and CO\u2082 but leave "
        "humidity, volatile organic compounds (VOCs), ambient light, electromagnetic "
        "fields, and barometric pressure entirely unmonitored."
    )

    add_body_paragraph(doc,
        "In this Perspective, we draw on evidence from the in vitro fertilization (IVF) "
        "field, molecular sensitivity data for PSCs, and a natural experiment in public "
        "databases to argue that environmental variables constitute an overlooked, "
        "correctable source of differentiation variability. We propose IoT-based "
        "environmental monitoring as a practical path forward."
    )

    # --- IVF precedent ---
    add_heading(doc, "The IVF Precedent: Environmental Control Transforms Outcomes",
                level=2)

    add_body_paragraph(doc,
        "The IVF field provides a compelling precedent. Before systematic environmental "
        "control, unexplained clinic-to-clinic variation was attributed to operator skill "
        "and patient selection. After implementing VOC filtration and positive-pressure "
        "air handling,{5} many laboratories saw 10\u201320% improvements in blastocyst "
        "rates.{6} Cai et al demonstrated through real-time monitoring that even low VOC "
        "concentrations linearly predicted decreased blastocyst quality, and that the "
        "effect was specific to fresh embryos\u2014frozen embryos were unaffected\u2014indicating "
        "that early-stage cells are selectively vulnerable to environmental perturbation.{7}"
    )

    add_body_paragraph(doc,
        "Seasonal patterns reinforce the environmental link. Leathersich et al analyzed "
        "3,659 frozen embryo transfer cycles in Perth, Australia, finding that oocytes "
        "collected in summer had 30% higher live birth rates than autumn-collected "
        "oocytes.{8} Crucially, only the season at oocyte collection mattered\u2014not the "
        "season of transfer\u2014establishing that the effect acts during a specific "
        "developmental window (oocyte maturation and early cleavage) rather than through "
        "a general mechanism. Northern Hemisphere studies report concordant summer "
        "advantages,{9} and direct experimental evidence in mice shows that embryos "
        "cultured during summer exhibit the 'two-cell block' phenomenon absent in "
        "winter.{10}"
    )

    add_body_paragraph(doc,
        "The IVF evidence provides a causal structure directly relevant to PSC biology. "
        "PSC differentiation protocols recapitulate aspects of early embryonic lineage "
        "commitment over 7\u201330 days and traverse epigenetic states homologous to "
        "preimplantation development\u2014exit from naive pluripotency, lineage priming, "
        "specification\u2014each potentially gated by the same environmental sensitivities "
        "that affect early embryos."
    )

    # --- PSC vulnerability ---
    add_heading(doc, "Why PSCs Are Uniquely Vulnerable", level=2)

    add_body_paragraph(doc,
        "Several lines of evidence support the premise that environmental factors affect "
        "PSCs while being negligible for established cell lines. First, PSCs display "
        "markedly greater chemical sensitivity: human iPS cells are approximately 1.5-fold "
        "more sensitive to drug exposure than ES cells, and both are several-fold less "
        "resistant than differentiated cell types; critically, naive-state induction "
        "further increases sensitivity, establishing a gradient from differentiated "
        "(resistant) to naive-pluripotent (most vulnerable).{11}"
    )

    add_body_paragraph(doc,
        "Second, osmotic perturbations of surprisingly small magnitude gate cell fate "
        "transitions. McCreery et al showed that compaction-triggered changes in nuclear "
        "shape and volume in human iPSCs remodel chromatin architecture and prime "
        "ectodermal differentiation.{12} During biosafety cabinet work, culture medium "
        "evaporates at rates determined by ambient humidity, producing osmolarity increases "
        "that could cross these thresholds. Chui et al confirmed that hyperosmotic culture "
        "drives NF-\u03baB and WNT signaling-dependent cell cycle exit and maturation{13}\u2014"
        "meaning osmolarity changes of physiological magnitude are not merely stress "
        "events but active differentiation signals."
    )

    add_body_paragraph(doc,
        "Third, PSCs depend on circadian clock components in a non-canonical manner. CRY1 "
        "is dramatically upregulated in iPSCs and ESCs compared to somatic cells, and its "
        "deletion impairs self-renewal and disrupts differentiation.{14} BMAL1 coordinates "
        "energy metabolism and differentiation independently of circadian oscillation, "
        "which is suppressed in pluripotent cells.{15} This architecture\u2014dependence on "
        "clock molecules without functional clock oscillation\u2014may render PSCs vulnerable "
        "to external zeitgeber inputs (eg, ambient light leaking into incubators or "
        "biosafety cabinets) that would not perturb differentiated cells with intact "
        "circadian buffering. Fourth, fine particulate matter downregulates pluripotency "
        "markers (NANOG, OCT4) in human ESCs through ROS-mediated AKT/ERK signaling{16} "
        "at concentrations harmless to somatic cells."
    )

    add_body_paragraph(doc,
        "Together, these findings establish that PSCs sit at an energetically shallow "
        "decision landscape where subtle environmental shifts can tip the balance between "
        "self-renewal and lineage commitment. Each mechanism operates through a different "
        "signal transduction pathway (osmolarity via nuclear mechanics, light via clock "
        "gene regulation, VOCs via ROS signaling), yet all converge on pluripotency gene "
        "networks. The key point is not that any single perturbation is catastrophic, "
        "but that the magnitude of real-world effects remains unknown because the "
        "relevant variables have not been systematically measured alongside "
        "differentiation outcomes."
    )

    # --- GEO natural experiment ---
    add_heading(doc, "A Natural Experiment Reveals Institutional Confounding", level=2)

    add_body_paragraph(doc,
        "To test whether institutional rhythms confound clonal research data, we analyzed "
        "submission dates for 6,101 PSC differentiation datasets deposited in NCBI GEO{17} "
        "between 2001 and 2026 (search: 'iPSC OR ESC differentiation'; inclusion: all "
        "dataset types with valid submission dates). The overall monthly distribution was "
        "significantly non-uniform (\u03c7\u00b2=32.3, df=11, P=.0007) with a pronounced "
        "March peak\u2014superficially consistent with a spring photoperiod advantage."
    )

    add_body_paragraph(doc,
        "We exploited a natural experiment: geographic variation in academic/fiscal year "
        "calendars. Countries were grouped by academic year start: Japan/South Korea "
        "(March/April, n=250), USA/Europe/China (September/October, n=2,825), and "
        "Southern Hemisphere (n=83). Using Pearson \u03c7\u00b2 and circular Rayleigh tests,{18} "
        "only the Japan/Korea group showed significant seasonality (\u03c7\u00b2=32.8, P=.0006, "
        "Rayleigh R=0.172) with a March peak perfectly aligned with Japan's fiscal "
        "year-end. The much larger September-start group showed no significant seasonal "
        "pattern (\u03c7\u00b2=15.0, P=.13). Japan\u2014just 3.3% of total datasets\u2014drove the "
        "overall signal entirely: excluding Japan eliminated statistical significance "
        "(Figure 1)."
    )

    # -- Insert Figure 1 inline after first mention --
    fig1_path = os.path.join(OUTPUT_DIR, "StemCells_Figure1.png")
    insert_figure(doc, fig1_path,
        "Institutional calendar confounding in GEO PSC datasets. "
        "(A) Monthly submission distributions for 6,101 GEO PSC differentiation datasets "
        "grouped by country academic-year start. (B) Country-level heatmap. "
        "(C) Circular Rayleigh vectors showing mean submission direction. "
        "(D) Statistical summary: only Japan/Korea shows significant seasonality, "
        "aligned with Japan's March 31 fiscal year-end.",
        fig_number=1
    )

    add_body_paragraph(doc,
        "This result carries a critical methodological lesson: what appeared to be "
        "evidence of seasonal biological influence on stem cell research was an artifact "
        "of Japan's March 31 fiscal year-end, which creates a rush to deposit data before "
        "annual reporting deadlines. The broader implication is that any field relying on "
        "publication or deposition timing as a proxy for experimental timing risks "
        "confounding institutional rhythms with biological ones. Prospective environmental "
        "measurement, not retrospective database mining, is required to detect real "
        "environmental effects on PSC differentiation."
    )

    # --- Monitoring proposal ---
    add_heading(doc, "The Path Forward: Environmental Monitoring as Method", level=2)

    add_body_paragraph(doc,
        "Our analysis demonstrates that retrospective database mining cannot detect or "
        "rule out environmental effects on PSC differentiation. The IVF evidence suggests "
        "such effects exist in early-stage cells,{8,10} and the vulnerability evidence "
        "suggests PSCs are plausible targets.{11,12} What is needed is a systematic "
        "method: prospective, controlled environmental monitoring paired with "
        "differentiation outcome tracking.{19}"
    )

    add_body_paragraph(doc,
        "We propose a three-phase approach (Figure 2). Phase I: equip PSC facilities with "
        "continuous multi-parameter sensors (humidity, illuminance, ELF-EMF, barometric "
        "pressure, VOCs, vibration) at 1-minute resolution alongside routine "
        "differentiation outcomes for 12 or more months. Modern IoT sensor packages "
        "achieve this at modest cost (<$5,000 per laboratory). Phase II: analyze existing "
        "IVF registry data (HFEA in the UK, ANZARD in Australia) containing date-stamped "
        "cycle outcomes with latitude information to directly test hemisphere-inversion "
        "predictions for photoperiod-driven effects. Phase III: for variables identified "
        "as significantly correlated with outcomes, conduct controlled interventions (eg, "
        "humidity-controlled vs uncontrolled incubators, light-tight hoods vs standard "
        "biosafety cabinets)."
    )

    # -- Insert Figure 2 inline after first mention --
    fig2_path = os.path.join(OUTPUT_DIR, "StemCells_Figure2.png")
    insert_figure(doc, fig2_path,
        "Proposed research roadmap for environmental profiling of PSC culture. "
        "(A) Evidence matrix classifying uncontrolled environmental variables by "
        "hemisphere-dependence and strength of existing evidence. (B) Three-phase "
        "investigation strategy: Phase I (passive IoT monitoring), Phase II "
        "(IVF registry analysis), Phase III (controlled intervention).",
        fig_number=2
    )

    add_body_paragraph(doc,
        "The hemisphere-inversion test provides a powerful diagnostic criterion. Variables "
        "whose effects are phase-inverted between Northern and Southern Hemisphere "
        "laboratories implicate solar-driven mechanisms\u2014photoperiod, temperature, UV "
        "flux\u2014while hemisphere-synchronous effects implicate geomagnetic or cosmic ray "
        "pathways. A coordinated multi-center study spanning both hemispheres could "
        "decompose seasonal variation into these mechanistic categories."
    )

    add_body_paragraph(doc,
        "A critical advantage of this approach is that it does not require investigators "
        "to alter their existing protocols. Environmental sensors operate passively in the "
        "background, producing time-series data that can be retrospectively correlated "
        "with differentiation outcomes already being collected. This dramatically lowers "
        "the barrier to participation and enables meta-analysis across laboratories using "
        "different protocols, cell lines, and differentiation targets."
    )

    # --- Implications ---
    add_heading(doc, "Implications for Reproducibility and Manufacturing", level=2)

    add_body_paragraph(doc,
        "As PSC-derived therapies advance toward clinical application,{1} manufacturing "
        "consistency becomes a regulatory imperative. Current Good Manufacturing Practice "
        "(cGMP) facilities control temperature and particulates, but do not routinely "
        "monitor illuminance, electromagnetic fields, VOCs, or barometric pressure. If "
        "any of these variables meaningfully affect differentiation efficiency\u2014as the "
        "IVF precedent suggests they affect embryo development\u2014current manufacturing "
        "harbors a hidden source of batch failure attributed to 'biological stochasticity' "
        "but potentially environmental in origin and therefore correctable.{20}"
    )

    add_body_paragraph(doc,
        "The economic argument is compelling. A single failed GMP differentiation batch "
        "can cost $50,000\u2013$500,000 in reagents and labor; environmental monitoring "
        "systems cost <$5,000 per facility to install. If monitoring identifies even one "
        "correctable environmental variable that reduces batch failure rates by 5\u201310%, "
        "the return on investment is immediate. Recent proposals for 'cytocentric' quality "
        "principles independently converge on this conclusion, arguing that environmental "
        "sensing must become a foundational element of cell manufacturing quality "
        "systems.{21,22}"
    )

    add_body_paragraph(doc,
        "The monitoring approach imposes no change on existing workflows\u2014it simply makes "
        "the invisible visible. Unlike prospective intervention studies requiring protocol "
        "changes and ethics approval, passive environmental monitoring can be implemented "
        "immediately in any laboratory or manufacturing facility. The resulting "
        "time-stamped, multi-parameter datasets, combined across sites and latitudes, "
        "would constitute a resource analogous to weather station networks that "
        "transformed meteorology from anecdote to predictive science. Even a null result "
        "would be informative: demonstrating that environmental variables do not "
        "meaningfully affect differentiation would definitively close this gap in our "
        "understanding and redirect attention to other sources of variability."
    )

    # --- Conclusion ---
    add_heading(doc, "Conclusion", level=2)

    add_body_paragraph(doc,
        "The history of experimental biology is punctuated by discoveries that 'noise' was "
        "signal from an unmeasured variable\u2014stomach ulcers attributed to stress were "
        "caused by Helicobacter pylori; sudden infant death syndrome attributed to fate "
        "was linked to sleeping position. We argue that inter-batch variability in PSC "
        "differentiation may represent a comparable situation: outcomes attributed to "
        "inherent biological stochasticity may partly reflect unmeasured environmental "
        "inputs."
    )

    add_body_paragraph(doc,
        "We have spent two decades optimizing recipes while ignoring the kitchen. The "
        "convergence of IVF environmental control evidence,{6,7} seasonal effects on "
        "early embryos,{8,10} molecular vulnerability of pluripotent cells,{11-16} and "
        "the methodological lessons from GEO analysis{17,18} suggest that the invisible "
        "variables are neither negligible nor intractable. They merely need to be "
        "measured. The cost of doing so is negligible compared to the cost of failed "
        "differentiation batches in clinical manufacturing. If even one invisible "
        "variable proves consequential\u2014as VOCs proved consequential in IVF{6}\u2014the "
        "return on investment will be transformative for both basic research "
        "reproducibility and clinical manufacturing."
    )

    # ──────────────────────────────────────────────
    # Acknowledgments
    # ──────────────────────────────────────────────
    add_heading(doc, "Acknowledgments", level=2)
    add_body_paragraph(doc, "[To be added]")

    # ──────────────────────────────────────────────
    # Author Contributions
    # ──────────────────────────────────────────────
    add_heading(doc, "Author Contributions", level=2)
    add_body_paragraph(doc, "[To be added upon submission]")

    # ──────────────────────────────────────────────
    # Disclosure of Potential Conflicts of Interest
    # ──────────────────────────────────────────────
    add_heading(doc, "Disclosure of Potential Conflicts of Interest", level=2)
    add_body_paragraph(doc, "The authors declare no potential conflicts of interest.")

    # ──────────────────────────────────────────────
    # Data Availability Statement
    # ──────────────────────────────────────────────
    add_heading(doc, "Data Availability Statement", level=2)
    add_body_paragraph(doc,
        "GEO dataset metadata and analysis code are available at "
        "[repository URL to be added upon acceptance]."
    )

    # ──────────────────────────────────────────────
    # Figure Legends
    # ──────────────────────────────────────────────
    add_heading(doc, "Figure Legends", level=2)

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0)
    run = p.add_run("Figure 1. ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run(
        "Academic calendar variation as a natural experiment. "
        "(A) Normalized monthly distributions of GEO PSC dataset submissions by academic "
        "year group. The March/April-start group (Japan/Korea) shows a pronounced March "
        "peak absent in the September-start group (USA/Europe/China). "
        "(B) Country-level heatmap showing that the pattern is driven predominantly by "
        "Japan. "
        "(C) Circular mean directions (Rayleigh vectors) for each group. "
        "(D) Statistical summary: only Japan/Korea achieves significance (P<.001). "
        "n=6,101 total datasets; 3,195 with country affiliation."
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    p = doc.add_paragraph()
    set_paragraph_format(p, line_spacing=2.0, space_before=Pt(12))
    run = p.add_run("Figure 2. ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run(
        "Research roadmap for environmental profiling of PSC culture. "
        "(A) Evidence matrix: uncontrolled environmental variables classified by "
        "hemisphere-dependence (inverted = solar-driven; synchronous = geomagnetic) and "
        "current evidence level. "
        "(B) Three-phase investigation strategy from passive monitoring through "
        "retrospective registry analysis to controlled intervention."
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # ──────────────────────────────────────────────
    # References (numbered, AMA style)
    # ──────────────────────────────────────────────
    add_heading(doc, "References", level=2)

    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, space_after=Pt(4), line_spacing=2.0)
        run = p.add_run(f"{i}. ")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run = p.add_run(ref)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    # Save
    out_path = os.path.join(OUTPUT_DIR, "StemCells_Perspective_InvisibleVariables.docx")
    doc.save(out_path)
    print(f"Manuscript saved to: {out_path}")

    # Word count (body text only, excluding abstract/significance/refs/legends)
    word_count = 0
    in_body = False
    skip_sections = {
        "Abstract", "Significance Statement", "Figure Legends",
        "References", "Acknowledgments", "Author Contributions",
        "Disclosure of Potential Conflicts of Interest",
        "Data Availability Statement"
    }
    current_section = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        # Track section headings
        if text in skip_sections:
            current_section = text
            in_body = False
            continue
        if text == "Introduction":
            in_body = True
            current_section = text
            continue
        # Other main body headings
        if text in (
            "The IVF Precedent: Environmental Control Transforms Outcomes",
            "Why PSCs Are Uniquely Vulnerable",
            "A Natural Experiment Reveals Institutional Confounding",
            "The Path Forward: Environmental Monitoring as Method",
            "Implications for Reproducibility and Manufacturing",
            "Conclusion"
        ):
            in_body = True
            current_section = text
            continue
        if in_body and text and current_section not in skip_sections:
            # Exclude inline figure captions from body word count
            if text.startswith("Figure ") and (". " in text[:12]):
                continue
            # Exclude figure placeholder text
            if text.startswith("[Figure ") and "placeholder" in text:
                continue
            word_count += len(text.split())
    print(f"Body word count (excl. abstract/significance/refs/legends): {word_count}")
    print(f"Target: <= 2,000 words")
    print(f"References: {len(REFERENCES)} (max 50)")
    print(f"Display items: 2 figures (max 2)")

    return out_path


if __name__ == "__main__":
    create_manuscript()
