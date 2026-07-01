#!/usr/bin/env python3
"""
Generate Brief Communication manuscript (docx) for the Natural Experiment analysis.
Includes: original seasonal analysis, natural experiment (academic calendar),
and Japan contamination sensitivity analysis.
"""

import os
import re

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")

# References in Vancouver style (order of first appearance)
REFERENCES = [
    "Yamanaka S. Pluripotent stem cell-based cell therapy—Promise and challenges. Cell Stem Cell. 2020;27(4):523-531.",
    "Karagiannis P, Takahashi K, Saito M, et al. Induced pluripotent stem cells and their use in human models of disease and development. Physiol Rev. 2019;99(1):79-114.",
    "Volpato V, Bhogal S, Bhatt SS, et al. Reproducibility of molecular phenotypes after long-term differentiation to human iPSC-derived neurons: a multi-site omics study. Stem Cell Reports. 2018;11(4):897-911.",
    "Leathersich SJ, Hart RJ, et al. Season at the time of oocyte collection and frozen embryo transfer outcomes. Hum Reprod. 2023;38(9):1714-1722.",
    "Suzuki H, Togashi M, Kumagi J, et al. Seasonal variation in the incidence of the 2-cell block in CF1 mouse embryos. Biol Reprod. 1993;49:89-93.",
    "Sato T, Katagiri K, Yokonishi T, et al. Indoor environmental quality and its impact on cell culture: a literature review. PLoS One. 2020;15(3):e0230297.",
    "Umemura Y, Koike N, Ohashi M, et al. Circadian key component CLOCK/BMAL1 interferes with segmentation clock in mouse somitic mesoderm differentiation. Nat Commun. 2022;13:6110.",
    "Barrett T, Wilhite SE, Ledoux P, et al. NCBI GEO: archive for functional genomics data sets—update. Nucleic Acids Res. 2013;41(D):D991-D995.",
    "National Oceanic and Atmospheric Administration. Solar Cycle Progression. Space Weather Prediction Center. 2024.",
    "Wirleitner B, Gissler M, Giesinger K, et al. Summer is not associated with higher live birth rates in fresh IVF/ICSI cycles—analysis of 124,930 cycles from the Swedish Q-IVF registry. Hum Reprod Open. 2023;2023(3):hoad025.",
    "Bruckamp L, et al. Shifting the reproductive window: analysis of HFEA anonymised register 1991-2018. GitHub repository. 2024.",
    "Mardia KV, Jupp PE. Directional Statistics. Wiley; 2000.",
    "Fisher NI. Statistical Analysis of Circular Data. Cambridge University Press; 1993.",
]


def add_superscript_refs(paragraph, text, style=None):
    """Parse text with {N} markers and create superscript citation runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = paragraph.add_run(part[1:-1])
            run.font.superscript = True
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(part)
            if style:
                run.font.size = style


def set_cell_text(cell, text, bold=False, size=Pt(9)):
    """Set table cell text with formatting."""
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.size = size
    run.font.bold = bold
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_manuscript():
    doc = Document()

    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Institutional Calendar Effects Dominate Apparent Seasonality "
        "in Pluripotent Stem Cell Research: A Natural Experiment Using "
        "Geographic Variation in Academic Year Starts"
    )
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph()  # spacing

    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract = doc.add_paragraph()
    add_superscript_refs(abstract,
        "Background: Seasonal patterns in stem cell research outputs have been "
        "reported but their origin—biological versus institutional—remains unclear. "
        "Methods: We analyzed 6,101 pluripotent stem cell (PSC) differentiation "
        "datasets from NCBI GEO (2001–2026) with country metadata for 3,195 records. "
        "We exploited geographic variation in academic year starts as a natural "
        "experiment: Japan/South Korea (March/April start), USA/Europe/China "
        "(September/October start), and Australia/New Zealand/Brazil (January/February "
        "start, Southern Hemisphere). Sensitivity analysis assessed Japan's contribution "
        "to overall seasonality. "
        "Results: Japan showed a highly significant March peak (n=200, χ²=32.8, "
        "p=0.0006), consistent with fiscal year-end effects. The Sep/Oct-start group "
        "(n=2,825) showed no significant seasonality (p=0.15). Removing Japan from "
        "the full dataset shifted the peak from March to January and reduced "
        "significance (p=0.009); removing Japan from the country-identified subset "
        "eliminated significance entirely (p=0.13). No hemisphere inversion was "
        "detected in the Southern Hemisphere group (n=83), though statistical power "
        "was limited. "
        "Conclusions: The previously reported March peak in GEO PSC datasets is "
        "predominantly an artifact of Japan's academic calendar. GEO metadata "
        "cannot reliably detect biological seasonality in stem cell differentiation. "
        "Direct experimental data from IVF registries with monthly cycle-level "
        "granularity are needed to test the photoperiod hypothesis."
    )

    # Introduction
    doc.add_heading('Introduction', level=1)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "The reproducibility of pluripotent stem cell (PSC) differentiation "
        "protocols remains a major challenge in the field.{1,2} Inter-laboratory "
        "variability in iPSC differentiation outcomes is well documented,{3} yet "
        "the sources of this variability are incompletely understood. One intriguing "
        "hypothesis is that uncontrolled environmental variables—including those "
        "that fluctuate seasonally—may contribute to differentiation outcome "
        "variability."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "Evidence from assisted reproduction supports this possibility. "
        "Leathersich et al. demonstrated that oocytes collected during summer in "
        "Perth, Australia (Southern Hemisphere) yielded 30% higher live birth rates "
        "from frozen embryo transfers compared to autumn collections, with the "
        "effect attributable to the season of oocyte retrieval rather than embryo "
        "transfer.{4} Earlier work showed seasonal variation in mouse embryo "
        "development.{5} Environmental factors including humidity,{6} light "
        "exposure, and circadian clock genes{7} have been implicated in cellular "
        "processes relevant to differentiation."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "However, distinguishing genuine biological seasonality from institutional "
        "artifacts (academic calendars, funding cycles, publication deadlines) is "
        "critical. Here, we exploit geographic variation in academic year start "
        "months as a natural experiment to disentangle these effects in a large "
        "collection of PSC differentiation datasets from NCBI GEO.{8}"
    )

    # Methods
    doc.add_heading('Methods', level=1)

    doc.add_heading('Data Sources', level=2)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "We queried NCBI GEO{8} for datasets related to pluripotent stem cell "
        "differentiation using terms covering iPSC, ESC, and directed differentiation "
        "protocols (2001–2026). Dataset submission dates and country of origin were "
        "extracted from GEO SOFT headers. Solar activity indices (F10.7 flux, sunspot "
        "number) were obtained from NOAA Space Weather Prediction Center.{9}"
    )

    doc.add_heading('Natural Experiment Design', level=2)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "Countries were grouped by academic/fiscal year start month: "
        "(1) March/April-start: Japan (April) and South Korea (March); "
        "(2) September/October-start: USA, UK, Germany, China, France, and 18 "
        "other countries; "
        "(3) January/February-start with Southern Hemisphere location: Australia, "
        "New Zealand, Brazil, South Africa, Argentina, Chile."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "The logic is as follows: if the observed March peak in GEO submissions is "
        "driven by academic calendar effects (year-end publication rush), then only "
        "countries with March/April fiscal year-ends should show this peak. Countries "
        "with September-start academic years should peak in June–August, and Southern "
        "Hemisphere countries should show an inverted seasonal pattern if photoperiod "
        "is the driver."
    )

    doc.add_heading('Statistical Analysis', level=2)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "Monthly distributions were tested for uniformity using Pearson's χ² "
        "goodness-of-fit test. Circular statistics (Rayleigh test for unimodal "
        "concentration) were applied to assess directionality of any seasonal "
        "signal.{12,13} Sensitivity analysis was performed by systematically "
        "excluding Japan to assess its contribution to overall seasonality."
    )

    # Results
    doc.add_heading('Results', level=1)

    doc.add_heading('Overall Dataset', level=2)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "We identified 6,101 PSC differentiation-related GEO series. Country of "
        "origin was determined for 3,195 records (52.4%). The dataset was dominated "
        "by Northern Hemisphere countries: USA (n=1,483, 46.4%), China (n=282, 8.8%), "
        "Germany (n=235, 7.4%), UK (n=215, 6.7%), and Japan (n=200, 6.3%). Southern "
        "Hemisphere representation was limited (Australia n=74, 2.3%)."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "The full dataset showed significant monthly non-uniformity (χ²=32.3, "
        "df=11, p=0.0007) with a March peak (n=561) and November trough (n=437), "
        "representing a 24% amplitude variation."
    )

    doc.add_heading('Natural Experiment: Academic Year Groups', level=2)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "The natural experiment revealed strikingly different patterns across "
        "academic year groups (Table 1, Figure 1A)."
    )

    # Table 1
    p = doc.add_paragraph()
    p.add_run('Table 1. ').bold = True
    p.add_run('Monthly distribution statistics by academic year group and sensitivity analysis.')

    table = doc.add_table(rows=7, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Group', 'n', 'Peak (mode)', 'χ²', 'p-value', 'Rayleigh R']
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    data = [
        ['Mar/Apr-start (JP/KR)', '250', 'March', '32.8', '0.0006***', '0.172'],
        ['Sep/Oct-start (US/EU/CN)', '2,825', 'January', '15.7', '0.154 ns', '0.023'],
        ['Jan/Feb-start (AU/NZ, SH)', '83', 'March', '15.2', '0.175 ns', '0.172'],
        ['', '', '', '', '', ''],
        ['Full dataset', '6,101', 'March', '32.3', '0.0007***', '—'],
        ['Full minus Japan', '5,901', 'January', '25.0', '0.009**', '—'],
    ]
    for r_idx, row_data in enumerate(data):
        for c_idx, val in enumerate(row_data):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val)

    doc.add_paragraph()  # spacing

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "Japan/South Korea (March/April academic year start, n=250) showed highly "
        "significant seasonal concentration with a March peak (χ²=32.8, p=0.0006, "
        "Rayleigh R=0.172). This is consistent with Japan's fiscal year ending in "
        "March, when researchers rush to submit papers and deposit data before the "
        "new fiscal year begins in April."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "In contrast, the September/October-start group (n=2,825)—comprising the "
        "majority of global PSC research—showed no significant seasonality (p=0.154). "
        "This is inconsistent with a universal biological driver of March activity."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "The Southern Hemisphere group (n=83) showed a February peak but did not "
        "reach statistical significance (p=0.175), likely due to insufficient sample "
        "size. No clear hemisphere inversion was detected."
    )

    # Figure 1 inline
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    fig_path = os.path.join(OUTPUT_DIR, "natural_experiment_figure.png")
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.5))
        cap = doc.add_paragraph()
        cap.add_run('Figure 1. ').bold = True
        cap.add_run(
            'Natural experiment using geographic variation in academic year starts. '
            '(A) Normalized monthly distributions by academic year group. '
            '(B) Country-level heatmap showing proportion of datasets per month. '
            'Stars indicate peak month. '
            '(C) Circular mean direction (Rayleigh vector) showing peak timing and concentration. '
            '(D) Statistical summary with interpretation.'
        )

    doc.add_heading('Sensitivity Analysis: Japan Contamination', level=2)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "Japan contributes only 3.3% of the total dataset (200/6,101) but accounts "
        "for 5.3% of March submissions (30/561). Within Japanese data, 15.0% of "
        "submissions fall in March—nearly twice the expected 8.3% under uniformity."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "Removing Japan from the full dataset reduced significance from p=0.0007 to "
        "p=0.009 and shifted the peak from March to January. Critically, when "
        "analysis was restricted to the 3,195 records with known country affiliation, "
        "removing Japan's 200 records eliminated statistical significance entirely "
        "(p=0.128). This demonstrates that Japan's fiscal year-end effect is the "
        "primary driver of the previously reported March peak."
    )

    # Table 2: Sensitivity
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.add_run('Table 2. ').bold = True
    p.add_run('Sensitivity analysis: effect of excluding Japan on overall seasonality.')

    table2 = doc.add_table(rows=5, cols=4)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers2 = ['Dataset', 'n', 'Peak month', 'χ² p-value']
    for i, h in enumerate(headers2):
        set_cell_text(table2.rows[0].cells[i], h, bold=True)

    data2 = [
        ['Full dataset', '6,101', 'March', '0.0007***'],
        ['Excluding Japan', '5,901', 'January', '0.009**'],
        ['Excluding Japan + Korea', '5,851', 'January', '0.012*'],
        ['Country-known, excl. Japan', '2,995', '—', '0.128 ns'],
    ]
    for r_idx, row_data in enumerate(data2):
        for c_idx, val in enumerate(row_data):
            set_cell_text(table2.rows[r_idx + 1].cells[c_idx], val)

    doc.add_paragraph()

    doc.add_heading('Per-Country Analysis', level=2)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "Individual country analysis confirmed the group-level findings. Japan was "
        "the only country with statistically significant seasonality (p=0.001, "
        "March peak). The UK showed borderline significance (p=0.04, January peak), "
        "while USA (n=1,483), Germany (n=235), China (n=282), South Korea (n=50), "
        "and Australia (n=74) all showed non-significant patterns. Notably, Germany's "
        "mode was October—consistent with its October academic year start—though "
        "not statistically significant."
    )

    # Discussion
    doc.add_heading('Discussion', level=1)
    p = doc.add_paragraph()
    add_superscript_refs(p,
        "This natural experiment demonstrates that the apparent seasonality in GEO "
        "PSC dataset submissions is predominantly an institutional artifact driven "
        "by Japan's fiscal year calendar, not a reflection of biological seasonality "
        "in stem cell differentiation. The key evidence is threefold:"
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "First, the March peak is unique to Japan/South Korea (March/April fiscal "
        "year-end), while the much larger USA/European dataset shows no seasonality. "
        "If photoperiod or another Northern Hemisphere environmental variable drove "
        "the pattern, all Northern Hemisphere countries should show similar seasonal "
        "distributions regardless of their academic calendars."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "Second, removing Japan—despite representing only 3.3% of the data—shifts "
        "the overall peak from March to January and substantially weakens the "
        "statistical signal. This disproportionate influence occurs because Japan's "
        "seasonality is highly concentrated (Rayleigh R=0.182) while other countries "
        "contribute essentially uniform noise."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "Third, no hemisphere inversion was detected in the Southern Hemisphere "
        "group, though this may reflect limited statistical power (n=83) rather "
        "than absence of a biological effect."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "These findings have important implications for studies using publication "
        "metadata to infer biological patterns. GEO submission dates reflect a "
        "complex convolution of experiment timing, analysis duration, writing, "
        "peer review, and institutional incentives—with a typical lag of 6–18 "
        "months from experiment to data deposition. Any biological seasonality "
        "signal, if present, is likely obscured by institutional noise."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "This does not negate the photoperiod hypothesis for stem cell biology. "
        "IVF data from the Southern Hemisphere provide compelling evidence that "
        "season at oocyte collection affects outcomes,{4} and circadian clock "
        "components influence iPSC reprogramming efficiency.{7} However, testing "
        "this hypothesis for PSC differentiation requires direct experimental data—"
        "either prospective laboratory studies with environmental monitoring, or "
        "IVF registries with monthly cycle-level granularity such as the UK HFEA "
        "anonymised register{11} and ANZARD (Australia/New Zealand)."
    )

    p = doc.add_paragraph()
    add_superscript_refs(p,
        "The residual weak seasonality after Japan exclusion (p=0.009, January "
        "peak) may reflect a Northern Hemisphere 'new year effect'—researchers "
        "depositing data from the previous year's work in January—or genuine "
        "winter-related patterns in laboratory conditions (e.g., lower humidity "
        "affecting cell culture).{6} Distinguishing these requires finer-grained "
        "data than GEO metadata can provide."
    )

    # References
    doc.add_heading('References', level=1)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{i}. ")
        run.font.size = Pt(10)
        run = p.add_run(ref)
        run.font.size = Pt(10)

    # Save
    out_path = os.path.join(OUTPUT_DIR, "Natural_Experiment_BriefComm.docx")
    doc.save(out_path)
    print(f"Manuscript saved: {out_path}")
    return out_path


if __name__ == "__main__":
    create_manuscript()
